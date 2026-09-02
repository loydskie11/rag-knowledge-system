import os
import io
import json
import time
from datetime import datetime, timedelta
from typing import List, Optional

# Third-party imports
from dotenv import load_dotenv
from fastapi import FastAPI, Depends, HTTPException, status, File, UploadFile, Form, Body, Request, Response, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordRequestForm, OAuth2PasswordBearer
from sqlalchemy.orm import Session
from sqlalchemy import text, or_
from pydantic import BaseModel, EmailStr
from openai import OpenAI
from groq import Groq
from supabase import create_client, Client
import PyPDF2
import groq
import time
import re
from rate_limiter import limiter
from sanitizer import sanitize_user_input, check_prompt_injection
import traceback
import numpy as np
from PIL import Image as PILImage

# Lazy-load PaddleOCR to avoid blocking at startup
_ocr_instance = None

def get_ocr():
    global _ocr_instance
    if _ocr_instance is None:
        from paddleocr import PaddleOCR
        _ocr_instance = PaddleOCR(use_angle_cls=True, lang='en')
    return _ocr_instance

def run_ocr(ocr_instance, img_array):
    """Run OCR and extract text lines. Compatible with PaddleOCR 2.x."""
    extracted = ""
    result = ocr_instance.ocr(img_array, cls=True)
    if result and result[0]:
        for line in result[0]:
            if line and len(line) >= 2:
                extracted += line[1][0] + "\n"
    return extracted

# Minimum characters per page from PyPDF2 before we OCR that page.
# Pages with tables/images embedded as images usually yield < 50 chars.
OCR_FALLBACK_THRESHOLD = 50

def extract_pdf_text(contents: bytes) -> str:
    """
    Fast, robust PDF text extraction with PyMuPDF (fitz) and selective OCR fallback.
    Prevents long OCR delays on multi-page PDFs with mixed text/graphic pages.
    """
    import fitz

    extracted_text = ""
    pdf_document = fitz.open(stream=contents, filetype="pdf")

    # 1. Fast text extraction using PyMuPDF across all pages
    for page in pdf_document:
        text = page.get_text() or ""
        if text.strip():
            extracted_text += text + "\n"

    # 2. If document contains sufficient text (>100 chars), return immediately! (0.1s processing time)
    if len(extracted_text.strip()) > 100:
        pdf_document.close()
        return extracted_text

    # 3. Fallback: If entire document yields < 100 chars (scanned document), run OCR on scanned pages
    ocr_text = ""
    ocr_instance = None
    max_ocr_pages = min(len(pdf_document), 30)  # Safety cap at 30 pages for OCR scans

    for page_num in range(max_ocr_pages):
        page_text = pdf_document[page_num].get_text() or ""
        if len(page_text.strip()) < 30:
            if ocr_instance is None:
                ocr_instance = get_ocr()
            fitz_page = pdf_document.load_page(page_num)
            pix = fitz_page.get_pixmap(dpi=120)
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            if pix.n == 4:
                img_array = img_array[:, :, :3]
            page_ocr = run_ocr(ocr_instance, img_array)
            ocr_text += (page_ocr if page_ocr.strip() else page_text) + "\n"
        else:
            ocr_text += page_text + "\n"

    pdf_document.close()
    return ocr_text if ocr_text.strip() else extracted_text


def supabase_query_with_retry(query_fn, retries=3, delay=1):
    """Retry a Supabase query on connection drops."""
    for attempt in range(retries):
        try:
            return query_fn()
        except Exception as e:
            if attempt < retries - 1 and "disconnect" in str(e).lower():
                time.sleep(delay)
                continue
            raise

# Local module imports
import models
import schemas
import utils
import vector_store
from database import engine, get_db


load_dotenv()
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY")
AI_API_KEY = os.getenv("AI_API_KEY", "ollama")
AI_BASE_URL = os.getenv("AI_BASE_URL", "http://localhost:11434/v1")

# Initialize clients — single instances used throughout the file
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
# We use OpenAI client pointing to the local Ollama server (OpenAI compatible endpoint)
groq_client = OpenAI(api_key=AI_API_KEY, base_url=AI_BASE_URL)

SMTP_USER     = os.getenv("EMAIL_ADDRESS")
SMTP_PASSWORD = os.getenv("EMAIL_APP_PASSWORD")


# ─────────────────────────────────────────────────────────────────────────────
# NOTIFICATION HELPERS
# ─────────────────────────────────────────────────────────────────────────────

VALID_NOTIF_TYPES = {"info", "success", "warning", "error"}


def _send_notification(user_email: str, n_type: str, title: str, message: str) -> None:
    """
    Insert one notification row for a specific user.
    Silent on failure — never blocks the main request.
    """
    if n_type not in VALID_NOTIF_TYPES:
        n_type = "info"
    try:
        supabase.table("notifications").insert({
            "user_email": user_email,
            "type":       n_type,
            "title":      title,
            "message":    message,
            "is_read":    False,
        }).execute()
    except Exception as exc:
        print(f"[_send_notification] silent failure — {exc}")


def _notify_all_admins(db: Session, n_type: str, title: str, message: str) -> None:
    """
    Send the same notification to every active ADMIN in the system.
    Only ADMINs receive this — faculty and students are excluded.
    """
    try:
        admins = (
            db.query(models.User)
            .filter(models.User.role == "ADMIN", models.User.status != "Disabled")
            .all()
        )
        for admin in admins:
            _send_notification(admin.email, n_type, title, message)
    except Exception as exc:
        print(f"[_notify_all_admins] silent failure — {exc}")


def _notify_role(db: Session, role: str, n_type: str, title: str, message: str) -> None:
    """
    Send the same notification to every active user of a given role.
    role: 'STUDENT' | 'FACULTY' | 'ADMIN'
    """
    try:
        users = (
            db.query(models.User)
            .filter(models.User.role == role, models.User.status != "Disabled")
            .all()
        )
        for user in users:
            _send_notification(user.email, n_type, title, message)
    except Exception as exc:
        print(f"[_notify_role:{role}] silent failure — {exc}")


def _notify_office(db: Session, office_name: str, n_type: str, title: str, message: str) -> None:
    """
    Broadcasts a notification to every active user whose administrative_office or department matches office_name.
    """
    if not office_name:
        return
    try:
        office_users = (
            db.query(models.User)
            .filter(
                models.User.status != "Disabled",
                (
                    (models.User.administrative_office == office_name) |
                    (models.User.department == office_name)
                )
            )
            .all()
        )
        for u in office_users:
            if u.email:
                _send_notification(u.email, n_type, title, message)
    except Exception as exc:
        print(f"[_notify_office] silent failure — {exc}")


def _notify_non_admin(db: Session, n_type: str, title: str, message: str) -> None:
    """
    Send to all active FACULTY and STUDENT users — excludes ADMINs.
    Used for document and announcement notifications.
    """
    try:
        users = (
            db.query(models.User)
            .filter(
                models.User.role.in_(["FACULTY", "STUDENT"]),
                models.User.status != "Disabled"
            )
            .all()
        )
        for user in users:
            _send_notification(user.email, n_type, title, message)
    except Exception as exc:
        print(f"[_notify_non_admin] silent failure — {exc}")


# ─────────────────────────────────────────────────────────────────────────────
# PYDANTIC SCHEMAS (request / response models)
# ─────────────────────────────────────────────────────────────────────────────

class ChatHistoryItem(BaseModel):
    role: str # "user" or "assistant" / "ai"
    content: str

class QuestionRequest(BaseModel):
    question:   str
    user_email: Optional[str] = "guest@ctu.edu.ph"
    user_role:  Optional[str] = "STUDENT"
    history:    Optional[List[ChatHistoryItem]] = []

class FeedbackRequest(BaseModel):
    question:   str
    answer:     str
    is_helpful: bool

class UpdateDocumentRequest(BaseModel):
    old_name:         str
    new_name:         str
    category:         str
    office:           str
    version:          str
    effectivity_date: str

class AccreditationUploadRequest(BaseModel):
    document_name: str
    program:       str
    area_code:     str
    category:      str = "Accreditation Evidence"

class ForgotPasswordRequest(BaseModel):
    email: EmailStr

class UpdatePasswordRequest(BaseModel):
    email:        EmailStr
    new_password: str

class UserCreateRequest(BaseModel):
    full_name: str
    email:     str
    role:      str
    password:  str
    administrative_office: Optional[str] = None
    is_iqa_auditor: Optional[bool] = False

class UserUpdateRequest(BaseModel):
    administrative_office: Optional[str] = None
    is_iqa_auditor: Optional[bool] = False

class ChangePasswordRequest(BaseModel):
    email: str
    current_password: str
    new_password: str

class UpdateProfileRequest(BaseModel):
    email: str       # The current email (to find them in the database)
    new_email: str   # The new email they want to change to
    full_name: str
    program: str

class AccessLogRequest(BaseModel):
    document_name: str
    action_type:   str
    user_email:    str
    user_role:     str

# Notification schemas
class NotificationCreate(BaseModel):
    user_email: str
    type:       str
    title:      str
    message:    str

class NotificationOut(BaseModel):
    id:         int
    user_email: str
    type:       str
    title:      str
    message:    str
    is_read:    bool
    created_at: str

# ── FIXED BroadcastRequest ───────────────────────────────────────
# target_role = "ALL" | "STUDENT" | "FACULTY"
# Admin is excluded from mass broadcasts — they manage the system.
class BroadcastRequest(BaseModel):
    title:       str
    message:     str
    target_role: str = "ALL"   # default: send to everyone (non-admin)
    sender_email: str = ""


# ─────────────────────────────────────────────────────────────────────────────
# APP SETUP
# ─────────────────────────────────────────────────────────────────────────────

class SendOTPRequest(BaseModel):
    email: EmailStr

class VerifyOTPRequest(BaseModel):
    email: EmailStr
    otp_code: str

class AccreditationReviewRequest(BaseModel):
    document_name: str
    status: str # "Approved" or "Needs Revision"
    feedback: str = ""


# --- SETTINGS SCHEMAS ---
class SettingsSchema(BaseModel):
    platform_name: str
    campus: str
    admin_email: str
    jwt_expiration: int
    otp_expiration: int
    ai_model: str
    ai_temperature: float
    ai_system_prompt: str
    rag_max_chunks: int

models.Base.metadata.create_all(bind=engine)

try:
    with engine.connect() as conn:
        conn.execute(text("ALTER TABLE iso_requirements ADD COLUMN IF NOT EXISTS cycle_year VARCHAR(50) DEFAULT '2025 Surveillance';"))
        conn.execute(text("ALTER TABLE iqa_day_schedules ADD COLUMN IF NOT EXISTS cycle_year VARCHAR(50) DEFAULT '2025 Surveillance';"))
        conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS administrative_office VARCHAR(255);"))
        conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS is_iqa_auditor BOOLEAN DEFAULT FALSE;"))
        conn.execute(text("ALTER TABLE paper_trail_records ADD COLUMN IF NOT EXISTS origin_office VARCHAR(255);"))
        conn.execute(text("ALTER TABLE paper_trail_records ADD COLUMN IF NOT EXISTS origin_person VARCHAR(255);"))
        conn.execute(text("ALTER TABLE paper_trail_records ADD COLUMN IF NOT EXISTS current_location VARCHAR(255);"))
        conn.execute(text("ALTER TABLE paper_trail_records ADD COLUMN IF NOT EXISTS transaction_type VARCHAR(50) DEFAULT 'Submission';"))
        conn.execute(text("ALTER TABLE qms_action_plans ADD COLUMN IF NOT EXISTS assessment_date VARCHAR(50);"))
        conn.execute(text("ALTER TABLE qms_action_plans ADD COLUMN IF NOT EXISTS assessment_notes TEXT;"))
        conn.execute(text("ALTER TABLE program_accreditations ADD COLUMN IF NOT EXISTS active_areas TEXT DEFAULT 'Area I,Area II,Area III,Area IV,Area V,Area VI,Area VII,Area VIII,Area IX,Area X';"))
        
        # Check if aaccup_requirements is empty, then seed
        res_aaccup = conn.execute(text("SELECT COUNT(*) FROM aaccup_requirements")).scalar()
        if res_aaccup == 0:
            seed_data = [
                ("Area I", "Vision, Mission, Goals and Objectives", "Approved Board Resolution of VMGO"),
                ("Area I", "Vision, Mission, Goals and Objectives", "Dissemination Evidence (Photos, Memos)"),
                ("Area I", "Vision, Mission, Goals and Objectives", "Stakeholder Awareness Survey Rating"),
                ("Area II", "Faculty", "Faculty Manual"),
                ("Area II", "Faculty", "201 Files / Credentials of Faculty"),
                ("Area II", "Faculty", "Faculty Development Plan"),
                ("Area II", "Faculty", "Summary of Faculty Workload and Loading"),
                ("Area III", "Curriculum and Instruction", "CMO / Syllabi for all subjects"),
                ("Area III", "Curriculum and Instruction", "Curriculum Map"),
                ("Area III", "Curriculum and Instruction", "Sample Exams and Rubrics"),
                ("Area IV", "Support to Students", "Student Handbook"),
                ("Area IV", "Support to Students", "Guidance and Counseling Reports"),
                ("Area IV", "Support to Students", "Student Organization Activities"),
                ("Area V", "Research", "Institutional Research Agenda"),
                ("Area V", "Research", "Published Research Papers"),
                ("Area V", "Research", "Research Incentives Memo"),
                ("Area VI", "Extension and Community Involvement", "Extension Program Plan"),
                ("Area VI", "Extension and Community Involvement", "MOA/MOU with Partner Communities"),
                ("Area VI", "Extension and Community Involvement", "Impact Assessment Report"),
                ("Area VII", "Library", "Library Manual"),
                ("Area VII", "Library", "Inventory of Books and Journals"),
                ("Area VII", "Library", "Library Utilization Reports"),
                ("Area VIII", "Physical Plant and Facilities", "Campus Development Plan"),
                ("Area VIII", "Physical Plant and Facilities", "Building Permits and Fire Safety Certs"),
                ("Area VIII", "Physical Plant and Facilities", "Maintenance Logs"),
                ("Area IX", "Laboratories", "Laboratory Manuals"),
                ("Area IX", "Laboratories", "Inventory of Equipment"),
                ("Area IX", "Laboratories", "Safety and Hazard Protocols"),
                ("Area X", "Administration", "Organizational Chart"),
                ("Area X", "Administration", "Strategic Plan"),
                ("Area X", "Administration", "Financial Audit Reports"),
            ]
            for code, title, desc in seed_data:
                conn.execute(text("INSERT INTO aaccup_requirements (id, area_code, area_title, description, created_at) VALUES (gen_random_uuid(), :code, :title, :desc, now())"), {"code": code, "title": title, "desc": desc})
            conn.commit()
            import time
            time.sleep(0.2)  # Allow DB transaction to finalize before math routes run
        else:
            conn.commit()
except Exception as _mig_err:
    print(f"Startup Migration Error: {_mig_err}")

app = FastAPI(
    title="CTU Institutional Knowledge System API",
    description="Backend for the RAG-Powered Quality Assurance System"
)

# ─────────────────────────────────────────────────────────────────────────────
# SECURITY HARDENING: CORS & HEADERS MIDDLEWARE
# ─────────────────────────────────────────────────────────────────────────────
allowed_origins_env = os.getenv("ALLOWED_ORIGINS")
if allowed_origins_env:
    allowed_origins = [origin.strip() for origin in allowed_origins_env.split(",") if origin.strip()]
else:
    allowed_origins = [
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://localhost:3000",
        "http://127.0.0.1:3000"
    ]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    """Applies Security Headers to defeat Clickjacking, MIME sniffing, and XSS attacks."""
    response = await call_next(request)
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    return response


# ─────────────────────────────────────────────────────────────────────────────
# MODULAR ROUTERS
# ─────────────────────────────────────────────────────────────────────────────
from routers import health, survey, announcements

app.include_router(health.router)
app.include_router(survey.router)
app.include_router(announcements.router)


# ─────────────────────────────────────────────────────────────────────────────
# AUTHENTICATION DEPENDENCIES (JWT VALIDATION)
# ─────────────────────────────────────────────────────────────────────────────
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login", auto_error=False)

def _extract_token_from_request(request: Request, bearer_token: Optional[str] = None) -> Optional[str]:
    """Helper to extract token from Bearer header OR HttpOnly cookie."""
    if bearer_token and bearer_token.strip() not in ("null", "undefined", ""):
        return bearer_token.strip()
    auth_header = request.headers.get("Authorization")
    if auth_header and auth_header.startswith("Bearer "):
        parts = auth_header.split(" ")
        if len(parts) > 1:
            token_val = parts[1].strip()
            if token_val and token_val not in ("null", "undefined", ""):
                return token_val
    cookie_token = request.cookies.get("access_token")
    if cookie_token:
        cookie_clean = cookie_token.replace("Bearer ", "").strip()
        if cookie_clean and cookie_clean not in ("null", "undefined", ""):
            return cookie_clean
    return None

def get_current_user(
    request: Request,
    token: Optional[str] = Depends(oauth2_scheme),
    db: Session = Depends(get_db)
) -> models.User:
    """Verifies incoming JWT access token from Header or HttpOnly Cookie."""
    extracted_token = _extract_token_from_request(request, token)
    if not extracted_token:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Authentication required. Missing Bearer access token or auth cookie.",
            headers={"WWW-Authenticate": "Bearer"},
        )
    try:
        payload = utils.decode_access_token(extracted_token)
        email: str = payload.get("sub")
        if not email:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid token claims.",
                headers={"WWW-Authenticate": "Bearer"},
            )
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Access token has expired or is invalid.",
            headers={"WWW-Authenticate": "Bearer"},
        )

    user = db.query(models.User).filter(models.User.email == email).first()
    if not user:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="User account not found.")

    if user.status != "Active":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="User account is disabled.")

    return user

def get_optional_user(
    request: Request,
    token: Optional[str] = Depends(oauth2_scheme),
    db: Session = Depends(get_db)
) -> Optional[models.User]:
    """Optional user resolution supporting both Header and Cookie auth."""
    extracted_token = _extract_token_from_request(request, token)
    if not extracted_token:
        return None
    try:
        payload = utils.decode_access_token(extracted_token)
        email = payload.get("sub")
        if email:
            return db.query(models.User).filter(models.User.email == email).first()
    except Exception:
        pass
    return None

def get_current_admin(current_user: models.User = Depends(get_current_user)) -> models.User:
    """Enforces Admin role requirement for administrative routes."""
    if current_user.role.upper() != "ADMIN":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access restricted. Administrator privileges required."
        )
    return current_user

def get_current_faculty_or_admin(current_user: models.User = Depends(get_current_user)) -> models.User:
    """Enforces Faculty or Admin role requirement."""
    if current_user.role.upper() not in ["FACULTY", "ADMIN"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access restricted to Faculty and Administrator roles."
        )
    return current_user



# ─────────────────────────────────────────────────────────────────────────────
# HEALTH
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/")
def read_root():
    return {"message": "Welcome to the CTU Knowledge System API! The server is running perfectly."}

@app.get("/test-db")
def test_db_connection(db: Session = Depends(get_db)):
    try:
        users_count = db.query(models.User).count()
        return {"status": "Success", "message": f"Connected to Supabase! Current users: {users_count}"}
    except Exception as e:
        return {"status": "Failed", "error": str(e)}
    
@app.post("/auth/send-otp")
def send_otp(request: Request, req: SendOTPRequest, db: Session = Depends(get_db)):
    # Rate Limit Enforcement: Max 3 OTP requests per 5 minutes per IP
    limiter.check(request, key_name="send_otp", max_requests=3, window_seconds=300)

    # 1. Check if email is already taken
    existing_user = db.query(models.User).filter(models.User.email == req.email).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="This email is already registered.")

    # 2. Generate OTP and Expiration
    otp = utils.generate_otp()
    expiration = datetime.utcnow() + timedelta(minutes=10)

    # 3. Update or create OTP record
    record = db.query(models.OTPVerification).filter(models.OTPVerification.email == req.email).first()
    if record:
        record.otp_code = otp
        record.expires_at = expiration
    else:
        record = models.OTPVerification(email=req.email, otp_code=otp, expires_at=expiration)
        db.add(record)
    db.commit()

    # 4. Send Email
    try:
        utils.send_otp_email(req.email, otp)
        return {"message": "Verification code sent to email."}
    except Exception as e:
        print(f"SMTP Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to send verification email. Please try again.")

@app.post("/auth/verify-otp")
def verify_otp(req: VerifyOTPRequest, db: Session = Depends(get_db)):
    record = db.query(models.OTPVerification).filter(models.OTPVerification.email == req.email).first()
    
    if not record or record.otp_code != req.otp_code:
        raise HTTPException(status_code=400, detail="Invalid verification code.")
    
    if datetime.utcnow() > record.expires_at:
        raise HTTPException(status_code=400, detail="This verification code has expired. Please request a new one.")

    return {"message": "Email verified successfully!"}
    

# ─────────────────────────────────────────────────────────────────────────────
# AUTH — REGISTER
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/register", response_model=schemas.UserResponse, status_code=status.HTTP_201_CREATED)
def register_user(user: schemas.UserCreate, db: Session = Depends(get_db)):
    # 1. Duplicate check
    existing_user = db.query(models.User).filter(models.User.email == user.email).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")

    # 2. Hash password — students are auto-verified, faculty require admin approval
    hashed_pwd  = utils.hash_password(user.password)
    auto_verify = user.role.upper() == "STUDENT"

    # 3. Persist user
    new_user = models.User(
        email=user.email,
        full_name=user.full_name,
        hashed_password=hashed_pwd,
        role=user.role.upper(),
        department=user.department,
        administrative_office=user.administrative_office,
        is_iqa_auditor=user.is_iqa_auditor or False,
        is_verified=auto_verify
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    # 4. Student profile
    if user.role.upper() == "STUDENT":
        new_student_profile = models.StudentProfile(
            user_id=new_user.id,
            course=user.course,
            year_level=user.year
        )
        db.add(new_student_profile)
        db.commit()

    # 5. Audit log
    try:
        supabase.table("system_events_logs").insert({
            "user_email":  new_user.email,
            "event_type":  "Account Creation",
            "description": f"New {new_user.role} account registered: {new_user.full_name}"
        }).execute()
    except Exception as exc:
        print(f"[register] audit log failed: {exc}")

    # 6. Notifications — role-specific, correct recipients only
    role_upper = new_user.role  # already uppercased above

    if role_upper == "FACULTY":
        # ── Tell the registering faculty member their status ──
        _send_notification(
            user_email=new_user.email,
            n_type="info",
            title="Account Pending Verification",
            message=(
                f"Hi {new_user.full_name or 'there'}! Your faculty account has been registered "
                "and is awaiting administrator verification. You'll be notified once approved."
            ),
        )
        # ── Notify ONLY admins (NOT other faculty, NOT students) ──
        _notify_all_admins(
            db=db,
            n_type="warning",
            title="New Faculty Awaiting Verification",
            message=(
                f"{new_user.full_name} ({new_user.email}) has registered as Faculty "
                "and is pending your verification. Go to Users & Roles to accept or decline."
            ),
        )

    elif role_upper == "STUDENT":
        # ── Welcome the student ──
        _send_notification(
            user_email=new_user.email,
            n_type="success",
            title="Welcome to CTU Argao KMS!",
            message=(
                f"Hi {new_user.full_name or 'there'}! Your student account is active. "
                "You can now access the Knowledge Repository, Ask Policy, and more."
            ),
        )
        # ── Notify ONLY admins (NOT faculty, NOT other students) ──
        _notify_all_admins(
            db=db,
            n_type="info",
            title="New Student Registered",
            message=(
                f"{new_user.full_name} ({new_user.email}) has created a new student account."
            ),
        )

    return new_user


# ─────────────────────────────────────────────────────────────────────────────
# AUTH — LOGIN
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/login", response_model=schemas.Token)
def login_user(
    request: Request,
    response: Response,
    form_data: OAuth2PasswordRequestForm = Depends(),
    db: Session = Depends(get_db)
):
    # Anti-Brute-Force Rate Limiting: Max 5 login attempts per minute per IP
    limiter.check(request, key_name="login_attempt", max_requests=5, window_seconds=60)

    user = db.query(models.User).filter(models.User.email == form_data.username).first()

    if not user or not utils.verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )

    if user.status == "Disabled":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Your account has been disabled. Please contact the system administrator."
        )

    if not user.is_verified:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Your Faculty account is currently pending Admin verification."
        )

    access_token = utils.create_access_token(data={"sub": user.email, "role": user.role})

    # Set HttpOnly, SameSite cookie so browser handles authentication securely without JS access
    response.set_cookie(
        key="access_token",
        value=f"Bearer {access_token}",
        httponly=True,
        max_age=60 * 60 * 8, # 8 hours
        samesite="lax",
        secure=False  # Set to True when deploying under HTTPS
    )

    # Audit log
    try:
        supabase.table("system_events_logs").insert({
            "user_email":  user.email,
            "event_type":  "Authentication",
            "description": "User successfully logged in"
        }).execute()
    except Exception as exc:
        print(f"[login] audit log failed: {exc}")

    user_dept = user.department
    if user.role == "STUDENT" and user.student_profile:
        user_dept = user.student_profile.course
    if not user_dept:
        user_dept = "ADMIN" if user.role == "ADMIN" else "Unassigned"

    return {
        "access_token": access_token,
        "token_type":   "bearer",
        "full_name":    user.full_name or "CTU User",
        "email":        user.email,
        "role":         user.role,
        "department":   user_dept,
        "administrative_office": user.administrative_office or "",
        "is_iqa_auditor": bool(user.is_iqa_auditor),
    }


@app.post("/logout")
def logout_user(response: Response):
    """Clears HttpOnly auth cookie on logout."""
    response.delete_cookie(key="access_token", samesite="lax")
    return {"message": "Successfully logged out."}



# ─────────────────────────────────────────────────────────────────────────────
# PASSWORD RESET
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/send-reset-email")
def send_reset_email(request: ForgotPasswordRequest, db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.email == request.email).first()
    if not user:
        raise HTTPException(status_code=404, detail="Email not found")

    try:
        utils.send_forgot_password_email(user.email)
        return {"message": "Success! Please check your email for the reset link."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")


@app.post("/update-password")
def update_password(request: UpdatePasswordRequest, db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.email == request.email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found in records")

    try:
        user.hashed_password = utils.hash_password(request.new_password)
        db.add(user)
        db.commit()
        db.refresh(user)

        _send_notification(
            user_email=request.email,
            n_type="warning",
            title="Password Changed",
            message=(
                "Your account password was recently changed. "
                "If you did not do this, contact the administrator immediately."
            ),
        )

        return {"message": "Password updated successfully!"}
    except Exception as e:
        db.rollback()
        print(f"CRITICAL DB ERROR: {str(e)}")
        raise HTTPException(status_code=500, detail="Could not update database")


# ─────────────────────────────────────────────────────────────────────────────
# USER MANAGEMENT
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/users", response_model=List[schemas.UserResponse])
def get_all_users(
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    return db.query(models.User).all()


@app.put("/users/{user_id}/verify")
def verify_user(
    user_id: str,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    user.is_verified = True
    db.commit()

    # Notify the faculty member — their account is now active
    _send_notification(
        user_email=user.email,
        n_type="success",
        title="Account Verified — You Can Now Log In",
        message=(
            "Your faculty account has been verified by the administrator. "
            "You now have full access to the CTU Argao Knowledge Management System."
        ),
    )

    return {"message": "User approved successfully!"}


@app.delete("/users/{user_id}")
def delete_user(
    user_id: str,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    db.delete(user)
    db.commit()
    return {"message": "User rejected and removed successfully!"}


@app.put("/users/{user_id}/disable")
def disable_user(
    user_id: str,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    user.status = "Disabled"
    db.commit()

    _send_notification(
        user_email=user.email,
        n_type="error",
        title="Account Disabled",
        message=(
            "Your account has been disabled by the system administrator. "
            "Please contact admin@ctu.edu.ph for assistance."
        ),
    )

    return {"message": f"Account for {user.full_name or user.email} has been disabled!"}


@app.put("/users/{user_id}/enable")
def enable_user(
    user_id: str,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    user.status = "Active"
    db.commit()

    _send_notification(
        user_email=user.email,
        n_type="success",
        title="Account Re-enabled",
        message=(
            "Your account has been re-enabled by the administrator. "
            "You can now log in to the system."
        ),
    )

    return {"message": f"Account for {user.full_name or user.email} has been re-enabled!"}


@app.post("/users")
def create_user_admin(request: UserCreateRequest, db: Session = Depends(get_db)):
    existing_user = db.query(models.User).filter(models.User.email == request.email).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="User with that email already exists.")

    hashed_pwd = utils.hash_password(request.password)

    new_user = models.User(
        email=request.email,
        full_name=request.full_name,
        hashed_password=hashed_pwd,
        role=request.role.upper(),
        department="ADMIN" if request.role.upper() == "ADMIN" else "Unassigned",
        administrative_office=request.administrative_office,
        is_iqa_auditor=request.is_iqa_auditor or False,
        is_verified=True
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    if request.role.upper() == "STUDENT":
        new_student_profile = models.StudentProfile(
            user_id=new_user.id,
            course="BSIT",
            year_level=1
        )
        db.add(new_student_profile)
        db.commit()

    # Notify the new user their account is ready
    _send_notification(
        user_email=new_user.email,
        n_type="success",
        title="Account Created by Administrator",
        message=(
            f"Your {request.role.capitalize()} account was created by the Administrator. "
            "You can now log in to the CTU Argao Knowledge Management System."
        ),
    )

    return {"message": f"{request.role.capitalize()} {request.full_name} created successfully!"}


@app.put("/users/{user_id}/details")
def update_user_details(
    user_id: str,
    payload: UserUpdateRequest,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Updates user administrative office and IQA Auditor role designation."""
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    user.administrative_office = payload.administrative_office
    user.is_iqa_auditor = bool(payload.is_iqa_auditor)
    db.commit()
    db.refresh(user)
    return {"message": "User administrative details updated successfully!"}


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTS — UPLOAD
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/upload-document")
def upload_document(
    file:             UploadFile = File(...),
    name:             str = Form(...),
    category:         str = Form(...),
    office:           str = Form(...),
    version:          str = Form(...),
    effectivity_date: str = Form(...),
    uploaded_by:      str = Form(None),   # pass userEmail from the frontend
    db: Session = Depends(get_db),        # ← added so we can fan-out notifications
):
    contents       = file.file.read()
    extracted_text = ""

    try:
        filename_lower = file.filename.lower()

        if filename_lower.endswith(".pdf"):
            extracted_text = extract_pdf_text(contents)
        elif filename_lower.endswith((".png", ".jpg", ".jpeg")):
            img = PILImage.open(io.BytesIO(contents)).convert("RGB")
            img_array = np.array(img)
            ocr = get_ocr()
            extracted_text = run_ocr(ocr, img_array)
        elif filename_lower.endswith(".txt"):
            extracted_text = contents.decode("utf-8")
        elif filename_lower.endswith(".docx"):
            import docx
            doc = docx.Document(io.BytesIO(contents))
            extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, TXT, or Image.")

        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from document or image.")

        safe_filename   = file.filename.replace(" ", "_")
        unique_filename = f"{int(time.time())}_{safe_filename}"

        supabase.storage.from_("documents").upload(
            file=contents,
            path=unique_filename,
            file_options={"content-type": "application/pdf"}
        )
        public_url = supabase.storage.from_("documents").get_public_url(unique_filename)

        metadata = {
            "name":             name,
            "category":         category,
            "office":           office,
            "version":          version,
            "effectivity_date": effectivity_date,
            "upload_date":      datetime.now().isoformat(),
            "file_url":         public_url,
            "status":           "Active",
            "uploaded_by":      uploaded_by or "Unknown",
        }

        chunks_count = vector_store.add_to_vector_db(extracted_text, metadata)

        # ── Notify admins ──
        _notify_all_admins(
            db=db,
            n_type="info",
            title="New Document Added to Repository",
            message=f"'{name}' (v{version}) was uploaded under {office} — {category}.",
        )

        # ── Notify all faculty and students ──
        _notify_non_admin(
            db=db,
            n_type="info",
            title="New Document Available",
            message=(
                f"A new document '{name}' (v{version}) has been added to the "
                f"Knowledge Repository under {category}. Check it out now."
            ),
        )

        return {
            "message": f"Successfully processed {name}!",
            "details": f"Created {chunks_count} searchable vector chunks and saved file to Storage."
        }

    except Exception as e:
        print(f"Error in upload-document: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


@app.post("/upload-new-version")
def upload_new_version(
    file:                 UploadFile = File(...),
    old_document_name:    str = Form(...),
    new_version:          str = Form(...),
    new_effectivity_date: str = Form(...),
    uploaded_by:          str = Form(None),
    db: Session = Depends(get_db),
):
    try:
        res = supabase.table("document_sections").select("metadata").eq("metadata->>name", old_document_name).limit(1).execute()
        if not res.data:
            raise HTTPException(status_code=404, detail="Old document not found in database.")

        old_metadata = res.data[0]['metadata']
        category     = old_metadata.get("category", "Policy")
        office       = old_metadata.get("office",   "Academic Affairs")

        old_chunks_res = supabase.table("document_sections").select("id, metadata").eq("metadata->>name", old_document_name).execute()
        if old_chunks_res.data:
            for chunk in old_chunks_res.data:
                chunk_meta           = chunk['metadata']
                chunk_meta['status'] = "Archived"
                supabase.table("document_sections").update({"metadata": chunk_meta}).eq("id", chunk['id']).execute()

        contents       = file.file.read()
        extracted_text = ""
        filename_lower = file.filename.lower()

        if filename_lower.endswith(".pdf"):
            extracted_text = extract_pdf_text(contents)
        elif filename_lower.endswith((".png", ".jpg", ".jpeg")):
            img = PILImage.open(io.BytesIO(contents)).convert("RGB")
            img_array = np.array(img)
            ocr = get_ocr()
            extracted_text = run_ocr(ocr, img_array)
        elif filename_lower.endswith(".txt"):
            extracted_text = contents.decode("utf-8")
        elif filename_lower.endswith(".docx"):
            import docx
            doc = docx.Document(io.BytesIO(contents))
            extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format.")

        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text.")

        safe_filename   = file.filename.replace(" ", "_")
        unique_filename = f"v{new_version}_{int(time.time())}_{safe_filename}"

        supabase.storage.from_("documents").upload(
            file=contents,
            path=unique_filename,
            file_options={"content-type": file.content_type}
        )
        public_url = supabase.storage.from_("documents").get_public_url(unique_filename)

        new_metadata = {
            "name":             old_document_name,
            "category":         category,
            "office":           office,
            "version":          new_version,
            "effectivity_date": new_effectivity_date,
            "upload_date":      datetime.now().isoformat(),
            "file_url":         public_url,
            "status":           "Active",
            "uploaded_by":      uploaded_by or "Unknown",
        }

        vector_store.add_to_vector_db(extracted_text, new_metadata)

        # ── Notify admins ──
        _notify_all_admins(
            db=db,
            n_type="info",
            title="Document Updated to New Version",
            message=(
                f"'{old_document_name}' has been updated to v{new_version}. "
                "The previous version has been archived automatically."
            ),
        )

        # ── Notify all faculty and students ──
        _notify_non_admin(
            db=db,
            n_type="info",
            title="Document Updated",
            message=(
                f"'{old_document_name}' has been updated to version {new_version}. "
                "The latest version is now available in the Knowledge Repository."
            ),
        )

        return {"message": "New version uploaded successfully and old version archived!"}

    except Exception as e:
        print(f"Update version error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTS — LIST
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/documents")
def get_documents():
    try:
        res = supabase.table("document_sections").select("metadata").execute()

        unique_docs = {}
        if res.data:
            for row in res.data:
                meta = row.get("metadata", {})
                if isinstance(meta, str):
                    try:    meta = json.loads(meta)
                    except: meta = {}

                name = meta.get("name")
                if not name or name in unique_docs:
                    continue

                unique_docs[name] = {
                    "name":             name,
                    "category":         meta.get("category",         ""),
                    "office":           meta.get("office",           ""),
                    "program":          meta.get("program",          "GLOBAL"),
                    "version":          meta.get("version",          "1.0"),
                    "effectivity_date": meta.get("effectivity_date", ""),
                    "status":           meta.get("status",           "Active"),
                    "file_url":         meta.get("file_url",         ""),
                    "upload_date":      meta.get("upload_date",      ""),
                    "uploaded_by":      meta.get("uploaded_by",      "Unknown"),
                }

        return list(unique_docs.values())

    except Exception as e:
        print(f"[get_documents] error: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch documents")



# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTS — UPDATE METADATA
# ─────────────────────────────────────────────────────────────────────────────

@app.put("/documents/update")
def update_document(request: UpdateDocumentRequest, db: Session = Depends(get_db)):
    try:
        res = supabase.table("document_sections").select("metadata").eq("metadata->>name", request.old_name).limit(1).execute()
        if not res.data:
            raise HTTPException(status_code=404, detail="Document not found")

        base_metadata = res.data[0]['metadata']

        base_metadata["name"]             = request.new_name
        base_metadata["category"]         = request.category
        base_metadata["office"]           = request.office
        base_metadata["version"]          = request.version
        base_metadata["effectivity_date"] = request.effectivity_date

        supabase.table("document_sections").update({"metadata": base_metadata}).eq("metadata->>name", request.old_name).execute()

        # ── Notify faculty and students of the edit ──
        if request.old_name != request.new_name:
            doc_label = f"'{request.old_name}' (renamed to '{request.new_name}')"
        else:
            doc_label = f"'{request.new_name}'"

        _notify_non_admin(
            db=db,
            n_type="info",
            title="Document Information Updated",
            message=(
                f"{doc_label} has been updated in the Knowledge Repository. "
                "The latest details are now available."
            ),
        )

        return {"message": "Document metadata updated successfully!"}
    except Exception as e:
        print(f"Update error: {e}")
        raise HTTPException(status_code=500, detail="Failed to update document")


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTS — ARCHIVE / DELETE
# ─────────────────────────────────────────────────────────────────────────────

@app.delete("/documents/{doc_name}")
def archive_document(doc_name: str, db: Session = Depends(get_db)):
    try:
        chunks_res = supabase.table("document_sections").select("id, metadata").eq("metadata->>name", doc_name).execute()

        if not chunks_res.data:
            raise HTTPException(status_code=404, detail="Document not found in the database.")

        for chunk in chunks_res.data:
            chunk_meta           = chunk['metadata']
            chunk_meta['status'] = "Archived"
            supabase.table("document_sections").update({"metadata": chunk_meta}).eq("id", chunk['id']).execute()

        # ── Notify faculty and students the document has been removed ──
        _notify_non_admin(
            db=db,
            n_type="warning",
            title="Document Removed from Repository",
            message=(
                f"'{doc_name}' has been archived and is no longer available "
                "in the Knowledge Repository."
            ),
        )

        return {"message": f"Document '{doc_name}' successfully archived and removed from active AI context!"}

    except Exception as e:
        print(f"Archive error: {e}")
        raise HTTPException(status_code=500, detail="Failed to archive document")


# ─────────────────────────────────────────────────────────────────────────────
# ASK POLICY
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/ask-policy")
def ask_policy(
    req: Request,
    request: QuestionRequest,
    db: Session = Depends(get_db),
    current_user: Optional[models.User] = Depends(get_optional_user)
):
    # RAG Query Rate Limiting: Max 15 questions per 1 minute per IP to prevent vector database scraping
    limiter.check(req, key_name="ask_policy", max_requests=15, window_seconds=60)

    raw_question = request.question or ""
    
    # SECURITY HARDENING: Sanitize input and detect prompt injection attempts
    question = sanitize_user_input(raw_question)
    is_injection, reason = check_prompt_injection(question)
    if is_injection:
        return {
            "answer": (
                "🛡️ **Security Guardrail Triggered**: Your query contains pattern signatures associated with "
                "prompt manipulation, system instruction override, or role bypass attempts. "
                "Please rephrase your question to ask directly about institutional policies."
            ),
            "sources": [],
            "restricted": False
        }

    # SECURITY HARDENING: Use cryptographically verified identity from JWT token if logged in
    if current_user:
        user_email = current_user.email
        user_role = current_user.role.upper()
    else:
        user_email = request.user_email
        user_role = (request.user_role or "STUDENT").upper()

    # ==========================================
    # ROLE-BASED ACCESS CONFIGURATION
    # ==========================================
    # Define which document categories are restricted per role.
    # STUDENT → cannot access Accreditation Evidence (confidential).
    # FACULTY and ADMIN → full access to all categories.
    STUDENT_RESTRICTED_CATEGORIES = ["Accreditation Evidence"]

    excluded_categories = []
    if user_role == "STUDENT":
        excluded_categories = STUDENT_RESTRICTED_CATEGORIES

    # Keywords that strongly indicate a query about restricted content.
    # Used to detect when a student is asking about a restricted topic
    # so we can return a clear explanation instead of a generic 'not found'.
    RESTRICTED_TOPIC_KEYWORDS = [
        "accreditation", "aaccup", "accredit", "self-survey",
        "survey instrument", "accreditation evidence",
        "area i", "area ii", "area iii", "area iv", "area v",
        "area vi", "area vii", "area viii", "area ix", "area x",
        "accreditation standard", "accreditor"
    ]

    # ==========================================
    # 1. FETCH DYNAMIC SETTINGS
    # ==========================================
    settings = db.query(models.SystemSettings).filter(models.SystemSettings.id == 1).first()
    
    # Fallbacks in case settings aren't set yet or cloud model name stored
    raw_model = settings.ai_model if settings else "qwen2.5"
    ai_model = "qwen2.5" if raw_model in ["qwen-2.5-32b", "qwen-2.5-7b-instruct", "llama-3.1-8b-instant", "llama-3.3-70b-versatile", "openai/gpt-oss-120b", "qwen/qwen3-32b"] else raw_model
    ai_temp = settings.ai_temperature if settings else 0.3
    base_prompt = settings.ai_system_prompt if settings else "You are the friendly and professional AI Policy Assistant for Cebu Technological University (CTU) Argao Campus."

    # ==========================================
    # 2. FAST-PATH: GREETING DETECTOR
    # ==========================================
    q_lower = question.strip().lower()
    chat_pattern = r'^(hi|hello|hey|good morning|good afternoon|good evening|greetings)[.!?,]*$'
    
    if re.match(chat_pattern, q_lower):
        greeting_prompt = f"""{base_prompt}
        YOUR PERSONALITY: Warm, welcoming, and helpful.
        The user just said: "{question}"
        Respond with a brief, warm greeting and ask how you can help them with university policies today.
        
        FORMATTING RULE:
        You must separate your main answer from the follow-up questions using exactly this string: |FOLLOWUPS|
        Everything after |FOLLOWUPS| must be written from the STUDENT'S or USER'S point of view — questions
        THEY might type next to the assistant (e.g. "What is the deadline for adding subjects?").
        NEVER use this section to ask the user a clarifying question yourself. If you genuinely need more
        information to answer well, put that clarifying question inside your main answer instead, and leave
        the |FOLLOWUPS| section empty.
        Put each follow-up question on a new line. Do not number them.
        """
        
        try:
            response = groq_client.chat.completions.create(
                messages=[{'role': 'system', 'content': greeting_prompt}],
                model=ai_model,
                temperature=0.5 # slightly higher for natural greetings
            )
            raw_answer = response.choices[0].message.content
            parts = raw_answer.split("|FOLLOWUPS|")
            answer = parts[0].strip()
            follow_ups = []
            if len(parts) > 1:
                follow_ups = [q.strip().lstrip('1234567890.- ') for q in parts[1].strip().split('\n') if q.strip()]
            
            # Save Chat History for Greeting
            from vector_store import supabase
            supabase.table("chat_history").insert({"user_email": user_email, "role": "user", "content": question}).execute()
            supabase.table("chat_history").insert({"user_email": user_email, "role": "ai", "content": answer}).execute()
            
            return {"answer": answer, "sources": [], "follow_ups": follow_ups[:3]}
        except Exception as e:
            print(f"Greeting Error: {e}")

    # ==========================================
    # 3. ORIGINAL TOPIC CATEGORIZATION
    # ==========================================
    try:
        from vector_store import supabase
        topic = "General"
        if len(question.split()) <= 2 or any(w in q_lower for w in ["test", "asdf"]):
            topic = "Ignored"
        elif "grade" in q_lower or "pass" in q_lower or "fail" in q_lower or "unit" in q_lower: topic = "Grading"
        elif "research" in q_lower or "publish" in q_lower or "incentive" in q_lower: topic = "Research"
        elif "faculty" in q_lower or "teacher" in q_lower or "leave" in q_lower: topic = "Faculty"
        elif "admission" in q_lower or "enroll" in q_lower or "shift" in q_lower: topic = "Admissions"
        elif "uniform" in q_lower or "dress" in q_lower or "id" in q_lower: topic = "Dress Code"

        supabase.table("query_logs").insert({
            "query_text": question,
            "topic_category": topic
        }).execute()
    except Exception as e:
        print(f"Failed to log query: {e}")

    history = request.history or []
    recent_history = [h for h in history if h.content and h.content.strip()][-6:]

    # ==========================================
    # 4. CONVERSATIONAL QUERY EXPANSION & SPELL CHECK
    # ==========================================
    if recent_history:
        history_summary = "\n".join([f"{'User' if h.role == 'user' else 'Assistant'}: {h.content}" for h in recent_history[-4:]])
        rewrite_prompt = f"""You are a search query optimizer for an institutional knowledge retrieval system.
Given the recent conversation history and a follow-up question, rephrase the follow-up question into a single, standalone, grammatically correct search query for a vector database.
Resolve any ambiguous pronouns or references (e.g. "it", "that", "the requirements", "the deadline", "how about that?") to the specific topic discussed.
Fix all typos.
Do NOT answer the question. ONLY output the optimized standalone search string.

Conversation History:
{history_summary}

Follow-up Question: "{question}"
Standalone Search Query:"""
    else:
        rewrite_prompt = f"""You are a search query optimizer. The user asked: "{question}"
Rewrite this into a clean, professional, grammatically correct search query optimized for a vector database. 
Fix all typos. Do not answer the question, ONLY output the optimized search string."""
    
    try:
        rewrite_response = groq_client.chat.completions.create(
            messages=[{'role': 'user', 'content': rewrite_prompt}],
            model=ai_model,
            temperature=0.0
        )
        optimized_query = rewrite_response.choices[0].message.content.strip().strip('"')
    except Exception:
        optimized_query = question

    # ==========================================
    # 5. ROLE-AWARE RAG RETRIEVAL
    # ==========================================
    # Pass excluded_categories so the vector search itself skips restricted
    # documents — we don't waste retrieval slots and never leak context to the AI.
    relevant_chunks = vector_store.search_knowledge(
        optimized_query, excluded_categories=excluded_categories
    )

    # Early exit: student is asking about a restricted topic.
    # Check this BEFORE the generic 'not found' so they get a clear reason.
    if user_role == "STUDENT" and excluded_categories:
        q_lower_check = (question + " " + optimized_query).strip().lower()
        is_restricted_query = any(kw in q_lower_check for kw in RESTRICTED_TOPIC_KEYWORDS)
        if is_restricted_query:
            return {
                "answer": "This information is restricted to faculty and administrators only. Accreditation documents are confidential and cannot be shared with students.",
                "sources": [],
                "follow_ups": [],
                "restricted": True
            }

    if not relevant_chunks:
        return {
            "answer": "I am sorry, but the specific document or policy regarding this matter is currently not available in our institutional knowledge repository.",
            "sources": [],
            "follow_ups": []
        }

    # Final safety net: strip any restricted chunks that slipped through
    # (should not happen after the retrieval-level filter, but belt-and-suspenders).
    safe_chunks = [
        chunk for chunk in relevant_chunks
        if chunk.get('metadata', {}).get('status') != 'Archived'
        and chunk.get('metadata', {}).get('category') not in excluded_categories
    ]
    relevant_chunks = safe_chunks
    context_text = "\n\n".join([chunk['content'] for chunk in relevant_chunks])

    # ==========================================
    # 6. ROLE-AWARE SYSTEM PROMPT (STRICTLY GROUNDED & ANTI-HALLUCINATION)
    # ==========================================
    role_context = {
        "STUDENT": (
            "The user is a STUDENT. Answer using verified public institutional documents "
            "such as the Student Handbook, academic catalog, grading guidelines, and general procedures. "
            "Accreditation evidence materials are confidential and must not be disclosed."
        ),
        "FACULTY": (
            "The user is a FACULTY MEMBER. Reference authorized institutional documents including "
            "the Faculty Manual, research policies, syllabus guidelines, and curriculum standards."
        ),
        "ADMIN": (
            "The user is an ADMINISTRATOR. You have access to all verified institutional policies "
            "and administrative quality standards."
        ),
    }.get(user_role, "The user's role is unknown. Answer conservatively using only general public policies.")

    system_prompt = f"""{base_prompt}
    
    You are the official CTU Argao Campus AI Policy Assistant. Your paramount duty is to provide 100% accurate, factual, and strictly grounded answers based EXCLUSIVELY on the verified institutional knowledge repository snippets provided below.

    USER CONTEXT:
    {role_context}

    STRICT FACTUAL INTEGRITY & ANTI-HALLUCINATION PROTOCOL:
    1. STRICT CONTEXT GROUNDING: You MUST base your entire response solely on the facts, requirements, and rules contained in the CONTEXT snippets below. DO NOT use pre-trained world knowledge, outside assumptions, or common-sense guesses to fill in gaps.
    2. ZERO FABRICATION POLICY: Never invent, extrapolate, or estimate deadlines, prerequisites, grade calculations, disciplinary penalties, fees, requirements, or staff names under any circumstances.
    3. MANDATORY REFUSAL WHEN UNCERTAIN / ABSENT: If the provided CONTEXT does not explicitly contain the necessary information to answer the user's question accurately, you MUST explicitly state:
       "I am sorry, but the specific document or policy regarding this matter is currently not available in our institutional knowledge repository."
       Do NOT attempt to give speculative general advice when the specific document is missing.
    4. CONVERSATIONAL TONE WITH STRICT FACTS: Be polite, warm, and helpful. You may structure the verified facts clearly using markdown headings, bullet points, or concise tables for readability, but every single fact must be directly traceable to the context.
    5. CITATION REFERENCE: Ground your answers clearly on the official document names and sections referenced in the context.

    FORMATTING RULE:
    You must separate your main answer from suggested follow-up questions using exactly this string: |FOLLOWUPS|
    Everything after |FOLLOWUPS| must be written from the STUDENT'S or USER'S point of view (e.g. "What is the deadline for adding subjects?").
    NEVER use this section to ask the user a clarifying question yourself. If you need more information, put it in your main answer and leave |FOLLOWUPS| empty.
    Put each follow-up question on a new line. Do not number them.
    
    CONTEXT FROM INSTITUTIONAL KNOWLEDGE REPOSITORY:
    {context_text}
    """

    # ==========================================
    # 7. AI CALL WITH MULTI-TURN CONVERSATION PAYLOAD
    # ==========================================
    messages_payload = [{'role': 'system', 'content': system_prompt}]
    
    # Append recent conversation history so the LLM understands conversation context
    for msg in recent_history:
        msg_role = 'user' if msg.role == 'user' else 'assistant'
        messages_payload.append({'role': msg_role, 'content': msg.content})
        
    # Append latest user question
    messages_payload.append({'role': 'user', 'content': question})

    # Grounded temperature: Keep low (max 0.15) to prevent creative hallucination
    grounded_temp = min(float(ai_temp), 0.15)

    try:
        response = groq_client.chat.completions.create(
            messages=messages_payload,
            model=ai_model,
            temperature=grounded_temp
        )

        raw_answer = response.choices[0].message.content
        parts = raw_answer.split("|FOLLOWUPS|")
        answer = parts[0].strip()

        follow_ups = []
        if len(parts) > 1:
            raw_questions = parts[1].strip().split('\n')
            for q in raw_questions:
                clean_q = q.strip().lstrip('1234567890.- ')
                if clean_q:
                    follow_ups.append(clean_q)
                if len(follow_ups) == 3: break

    except Exception as e:
        print(f"Cloud API Error: {e}")
        return {"answer": "I'm having a bit of trouble connecting to the network. Please try again in a moment!", "sources": [], "follow_ups": []}

    # ==========================================
    # 7. ORIGINAL SOURCE FORMATTING & MATH
    # ==========================================
    unique_sources = {}
    for chunk in relevant_chunks:
        source_name = f"{chunk['metadata']['name']} (v{chunk.get('metadata', {}).get('version', '1.0')}) - {chunk['metadata']['office']}"

        if source_name not in unique_sources:
            clean_snippet = chunk['content'][:150].strip() + "..."
            raw_similarity = chunk.get('similarity', 0.85)
            human_score = raw_similarity * 1.5 
            relevance_percentage = min(99, int(human_score * 100))

            unique_sources[source_name] = {
                "name": source_name,
                "snippet": clean_snippet,
                "relevance": relevance_percentage 
            }

    sources = list(unique_sources.values())
    final_sources = sources if "I'm sorry, I don't have" not in answer and len(context_text) > 10 else []

    # ==========================================
    # 8. ORIGINAL CHAT HISTORY SAVING
    # ==========================================
    try:
        supabase.table("chat_history").insert({"user_email": user_email, "role": "user", "content": question}).execute()
        supabase.table("chat_history").insert({"user_email": user_email, "role": "ai", "content": answer}).execute()
    except Exception as e:
        print(f"Failed to save chat history: {e}")

    return {
        "answer": answer,
        "sources": final_sources,
        "follow_ups": follow_ups
    }


# ─────────────────────────────────────────────────────────────────────────────
# BROADCAST ANNOUNCEMENT  ← FIXED
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/admin/broadcast")
def broadcast_notification(request: BroadcastRequest, db: Session = Depends(get_db)):
    """
    Send an announcement notification to users by role.
    target_role = "ALL"     → FACULTY + STUDENT (not admins)
    target_role = "STUDENT" → students only
    target_role = "FACULTY" → faculty only

    Admins are deliberately excluded — they are the senders, not receivers.
    The admin who sends it gets a confirmation notification instead.
    """
    try:
        sent = 0

        if request.target_role.upper() in ("ALL", "EVERYONE"):
            # Faculty + Students
            users = (
                db.query(models.User)
                .filter(
                    models.User.role.in_(["FACULTY", "STUDENT"]),
                    models.User.status != "Disabled"
                )
                .all()
            )
        else:
            users = (
                db.query(models.User)
                .filter(
                    models.User.role == request.target_role.upper(),
                    models.User.status != "Disabled"
                )
                .all()
            )

        for user in users:
            _send_notification(
                user_email=user.email,
                n_type="info",
                title=f"📢 {request.title}",
                message=request.message,
            )
            sent += 1

        # ── Confirm back to the sending admin ──
        if request.sender_email:
            _send_notification(
                user_email=request.sender_email,
                n_type="success",
                title="Announcement Sent",
                message=(
                    f"Your announcement '{request.title}' was successfully "
                    f"delivered to {sent} user(s)."
                ),
            )

        # ── Log the broadcast ──
        try:
            supabase.table("system_events_logs").insert({
                "user_email":  request.sender_email or "admin",
                "event_type":  "Broadcast Announcement",
                "description": f"Announcement '{request.title}' sent to {sent} {request.target_role} user(s)."
            }).execute()
        except Exception as exc:
            print(f"[broadcast] audit log failed: {exc}")

        return {
            "message": f"Broadcast sent to {sent} {request.target_role} user(s).",
            "sent":    sent
        }

    except Exception as e:
        print(f"[broadcast] error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────────────────────────────────────
# SYSTEM STATS
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/system-stats")
def get_system_stats(role: str = "STUDENT", db: Session = Depends(get_db)):
    import json

    total_docs     = 0
    total_users    = 0
    total_queries  = 0
    qms_total      = 0
    qms_overdue    = 0
    aaccup_pending = 0
    iso_pending    = 0
    iso_compliance = 0

    try:
        total_users = db.query(models.User).count()
    except Exception as e:
        print(f"User count error: {e}")

    try:
        docs_response  = supabase.table("document_sections").select("metadata").execute()
        unique_docs    = set()
        unique_pending = set()
        if docs_response.data:
            for row in docs_response.data:
                meta = row.get("metadata", {})

                if isinstance(meta, str):
                    try:    meta = json.loads(meta)
                    except: meta = {}

                doc_name = meta.get("document_name") or meta.get("title") or meta.get("name") or meta.get("file_name") or meta.get("source")
                category = meta.get("category", "")
                status   = meta.get("status", "Active")

                if doc_name:
                    if str(status).lower() == "archived":
                        continue
                    # Students only see public institutional policies/handbooks, not internal accreditation evidence
                    if role.upper() == "STUDENT" and category == "Accreditation Evidence":
                        continue
                    unique_docs.add(doc_name)

                if category == "Accreditation Evidence" and str(status).lower() == "pending" and doc_name:
                    unique_pending.add(doc_name)

        # For ADMIN: include total records across repository, paper trail, and QA evidences
        if role.upper() == "ADMIN":
            try:
                ched_ev = db.query(models.ChedEvidence.file_name).all()
                for (cf,) in ched_ev:
                    if cf: unique_docs.add(cf)
                
                iso_ev = db.query(models.ISOEvidence.file_name).all()
                for (isf,) in iso_ev:
                    if isf: unique_docs.add(isf)

                pt_records = db.query(models.PaperTrailRecord.title).all()
                for (pt,) in pt_records:
                    if pt: unique_docs.add(pt)

                qms_ev = db.query(models.QMSEvidence.file_name).all()
                for (qf,) in qms_ev:
                    if qf: unique_docs.add(qf)
            except Exception as ex:
                print(f"Admin document aggregate error: {ex}")

        total_docs = len(unique_docs)
        aaccup_pending = len(unique_pending)
    except Exception as e:
        print(f"Document count error: {e}")

    try:
        query_response = supabase.table("chat_history").select("id").eq("role", "user").execute()
        total_queries  = len(query_response.data) if query_response.data else 0
    except Exception as e:
        print(f"Chat count error: {e}")

    # QMS Action Plans (Form 6)
    try:
        qms_total = db.query(models.QMSActionPlan).count()
        qms_plans = db.query(models.QMSActionPlan).all()
        overdue_count = 0
        now_date_str = datetime.utcnow().strftime("%Y-%m-%d")
        for p in qms_plans:
            if p.status == "Overdue":
                overdue_count += 1
            elif p.status != "Completed" and p.target_date and str(p.target_date) < now_date_str:
                overdue_count += 1
        qms_overdue = overdue_count
    except Exception as e:
        print(f"QMS stats error: {e}")

    # ISO Requirements & Campus Compliance
    try:
        iso_pending = db.query(models.ISORequirement).filter(models.ISORequirement.status == "Pending").count()
        
        iso_reqs = db.query(models.ISORequirement).filter(models.ISORequirement.cycle_year.ilike("%2026%")).all()
        if not iso_reqs:
            iso_reqs = db.query(models.ISORequirement).all()
        
        total_iso = len(iso_reqs)
        compliant_iso = len([r for r in iso_reqs if r.status == "Compliant"])
        iso_compliance = int((compliant_iso / total_iso) * 100) if total_iso > 0 else 0
    except Exception as e:
        print(f"ISO stats error: {e}")

    return {
        "documents":      total_docs,
        "users":          total_users,
        "queries":        total_queries,
        "qms_total":      qms_total,
        "qms_overdue":    qms_overdue,
        "aaccup_pending": aaccup_pending,
        "iso_pending":    iso_pending,
        "iso_compliance": iso_compliance,
        "qmsTotal":       qms_total,
        "qmsOverdue":     qms_overdue,
        "aaccupPending":  aaccup_pending,
        "isoPending":     iso_pending,
        "isoCompliance":  iso_compliance,
    }


# ─────────────────────────────────────────────────────────────────────────────
# FEEDBACK
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/feedback")
def submit_feedback(request: FeedbackRequest):
    try:
        supabase.table("feedback_logs").insert({
            "question_text": request.question,
            "ai_response":   request.answer,
            "is_helpful":    request.is_helpful
        }).execute()
        return {"status": "success", "message": "Feedback recorded!"}
    except Exception as e:
        print(f"Failed to log feedback: {e}")
        return {"status": "error", "message": str(e)}


# ─────────────────────────────────────────────────────────────────────────────
# CHAT HISTORY
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/chat-history")
def get_chat_history(email: str):
    seven_days_ago = (datetime.now() - timedelta(days=7)).isoformat()

    try:
        response = supabase.table("chat_history")\
            .select("*")\
            .eq("user_email", email)\
            .gte("created_at", seven_days_ago)\
            .order("created_at", desc=False)\
            .execute()
        return response.data
    except Exception as e:
        print(f"Error fetching history: {e}")
        return []



# ─────────────────────────────────────────────────────────────────────────────
# ACCREDITATION
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/upload-accreditation-evidence")
async def upload_accreditation_evidence(
    file:                UploadFile = File(...),
    document_name:       str = Form(...),
    program:             str = Form(...),
    area_code:           str = Form(...),
    requirement_target:  str = Form(...),
    uploaded_by:         str = Form(None),
):
    import time, io, PyPDF2
    from datetime import datetime

    try:
        contents       = await file.read()
        extracted_text = ""
        filename_lower = file.filename.lower()

        if filename_lower.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(contents))
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text: extracted_text += text + "\n"
        elif filename_lower.endswith(".txt"):
            extracted_text = contents.decode("utf-8")
        elif filename_lower.endswith(".docx"):
            import docx
            doc = docx.Document(io.BytesIO(contents))
            extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format.")

        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text.")

        safe_filename   = file.filename.replace(" ", "_")
        unique_filename = f"evid_{int(time.time())}_{safe_filename}"

        supabase.storage.from_("documents").upload(
            file=contents, path=unique_filename, file_options={"content-type": file.content_type}
        )
        public_url = supabase.storage.from_("documents").get_public_url(unique_filename)

        metadata = {
            "name": document_name,
            "category": "Accreditation Evidence",
            "office": "Quality Assurance",
            "version": "1.0",
            "status": "Pending", 
            "program": program,
            "area_code": area_code,
            "requirement_target": requirement_target, 
            "uploaded_by": uploaded_by, 
            "admin_feedback": "",
            "upload_date": datetime.now().isoformat(),
            "file_url": public_url
        }

        vector_store.add_to_vector_db(extracted_text, metadata)

        # --- SILENT AUDIT LOG ---
        try:
            # FIX: Using the global supabase client instead of importing it locally
            supabase.table("system_events_logs").insert({
                "user_email": uploaded_by, 
                "event_type": "Accreditation Upload",
                "description": f"Uploaded '{document_name}' for {program} ({area_code}) - Pending Review"
            }).execute()
        except Exception as e:
            print(f"Failed to log accreditation upload: {e}")
        # ------------------------

        return {"message": "Evidence successfully uploaded and is pending Admin review!"}

    except Exception as e:
        print(f"Evidence upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- NEW: ADMIN REVIEW QUEUE ROUTE ---
@app.get("/admin/accreditation-pending")
def get_pending_accreditation():
    from vector_store import supabase
    try:
        res = supabase.table("document_sections").select("metadata").eq("metadata->>category", "Accreditation Evidence").eq("metadata->>status", "Pending").execute()
        
        unique_docs = {}
        if res.data:
            for item in res.data:
                meta = item.get('metadata', {})
                doc_name = meta.get('name')
                if doc_name and doc_name not in unique_docs:
                    unique_docs[doc_name] = {
                        "name": doc_name,
                        "program": meta.get('program', 'Unknown'),
                        "area_code": meta.get('area_code', 'Unknown'),
                        "target": meta.get('requirement_target', 'Unknown'),
                        "uploaded_by": meta.get('uploaded_by', 'Faculty Member'),
                        "date": meta.get('upload_date', '').split('T')[0],
                        "url": meta.get('file_url')
                    }
        return list(unique_docs.values())
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to fetch pending queue")

# --- NEW: ADMIN APPROVAL/REJECTION ROUTE ---
@app.post("/admin/accreditation-review")
def review_accreditation(req: AccreditationReviewRequest):
    from vector_store import supabase
    try:
        # Fetch all chunks of this document
        chunks_res = supabase.table("document_sections").select("id, metadata").eq("metadata->>name", req.document_name).execute()
        
        if not chunks_res.data:
            raise HTTPException(status_code=404, detail="Document not found")

        # Update status and feedback for all chunks
        for chunk in chunks_res.data:
            chunk_meta = chunk['metadata']
            chunk_meta['status'] = req.status
            chunk_meta['admin_feedback'] = req.feedback
            
            supabase.table("document_sections").update({"metadata": chunk_meta}).eq("id", chunk['id']).execute()

        # --- SILENT AUDIT LOG ---
        try:
            action_desc = f"Admin marked '{req.document_name}' as {req.status}."
            if req.status == "Needs Revision":
                action_desc += f" Feedback: {req.feedback}"
                
            supabase.table("system_events_logs").insert({
                "user_email": "System Admin", # Can be dynamic if you pass admin email
                "event_type": "QA Review",
                "description": action_desc
            }).execute()
        except Exception as e:
            print(f"Failed to log QA Review: {e}")
        # ------------------------

        return {"message": f"Document successfully marked as {req.status}!"}
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to process review")

# 2. THE UPDATED STATUS ROUTE (Database-driven AACCUP requirements)
@app.get("/accreditation-status/{program}")
def get_accreditation_status(program: str, db: Session = Depends(get_db)):
    try:
        # Fetch active areas configuration for the program if exists
        active_areas_set = None
        prog_acc = db.query(models.ProgramAccreditation).filter(models.ProgramAccreditation.program_code == program).first()
        if prog_acc and prog_acc.active_areas:
            active_areas_set = set(a.strip() for a in prog_acc.active_areas.split(",") if a.strip())

        res = supabase.table("document_sections").select("metadata").eq("metadata->>category", "Accreditation Evidence").eq("metadata->>program", program).eq("metadata->>status", "Approved").execute()
        
        fulfilled_reqs = {}
        if res.data:
            for item in res.data:
                meta = item.get('metadata', {})
                area = meta.get('area_code')
                req_target = meta.get('requirement_target')
                
                if area and req_target:
                    if area not in fulfilled_reqs:
                        fulfilled_reqs[area] = set()
                    fulfilled_reqs[area].add(req_target)

        # Query database requirements
        db_reqs = db.query(models.AaccupRequirement).all()
        area_groups = {}
        for r in db_reqs:
            if r.area_code not in area_groups:
                area_groups[r.area_code] = {"title": r.area_title, "reqs": []}
            area_groups[r.area_code]["reqs"].append(r.description)

        # Order areas logically Areas I to X
        sorted_codes = sorted(area_groups.keys())

        areas_list = []
        total_req  = 0
        total_ev   = 0

        for code in sorted_codes:
            if active_areas_set is not None and code not in active_areas_set:
                continue

            reqs = area_groups[code]["reqs"]
            title = area_groups[code]["title"]
            required_count = len(reqs)
            ev_count       = len(fulfilled_reqs.get(code, set()))
            capped_ev      = min(ev_count, required_count)
            total_req     += required_count
            total_ev      += capped_ev
            comp           = int((capped_ev / required_count) * 100) if required_count > 0 else 0

            areas_list.append({
                "id":            code.replace(" ", ""),
                "code":          code,
                "title":         title,
                "compliance":    comp,
                "required":      required_count,
                "evidenceCount": ev_count,
                "gaps":          max(0, required_count - ev_count),
            })

        overall = int((total_ev / total_req) * 100) if total_req > 0 else 0

        return {
            "level":    "Level III" if program == "BSIT" else "Level II",
            "overall":  overall,
            "gaps":     total_req - total_ev,
            "evidence": total_ev,
            "areas":    areas_list,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to calculate status")


@app.get("/accreditation-details/{program}/{area_code}")
def get_accreditation_details(program: str, area_code: str, db: Session = Depends(get_db)):
    try:
        response = supabase.table("document_sections")\
            .select("metadata")\
            .eq("metadata->>category", "Accreditation Evidence")\
            .eq("metadata->>program",  program)\
            .eq("metadata->>area_code", area_code)\
            .execute()

        unique_docs       = {}
        fulfilled_targets = set()

        if response.data:
            for item in response.data:
                meta = item.get('metadata', {})
                if meta.get('status') == "Archived": continue
                
                doc_name = meta.get('name')
                req_target = meta.get('requirement_target') 
                doc_status = meta.get('status', 'Pending')
                
                if req_target and doc_status == "Approved":
                    fulfilled_targets.add(req_target) 
                
                if doc_name and doc_name not in unique_docs:
                    unique_docs[doc_name] = {
                        "name": doc_name,
                        "date": meta.get('upload_date', '').split('T')[0] if 'upload_date' in meta else 'Recently',
                        "url": meta.get('file_url') or meta.get('source', '#'),
                        "target": req_target or "Uncategorized",
                        "status": doc_status,
                        "feedback": meta.get('admin_feedback', '')
                    }

        uploaded_files = list(unique_docs.values())

        db_reqs = db.query(models.AaccupRequirement).filter(models.AaccupRequirement.area_code == area_code).all()
        requirements  = []
        for req in db_reqs:
            requirements.append({
                "id":     str(req.id),
                "text":   req.description,
                "area_title": req.area_title,
                "is_met": req.description in fulfilled_targets,
            })

        return {
            "requirements":  requirements,
            "uploadedFiles": sorted(uploaded_files, key=lambda x: x['date'], reverse=True),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to fetch details")


# --- AACCUP REQUIREMENTS CRUD ENDPOINTS ---
@app.get("/aaccup/requirements", response_model=List[schemas.AaccupRequirementResponse])
def get_aaccup_requirements(area_code: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(models.AaccupRequirement)
    if area_code:
        query = query.filter(models.AaccupRequirement.area_code == area_code)
    return query.order_by(models.AaccupRequirement.area_code.asc(), models.AaccupRequirement.created_at.asc()).all()

@app.post("/aaccup/requirements", response_model=schemas.AaccupRequirementResponse)
def create_aaccup_requirement(req: schemas.AaccupRequirementCreate, db: Session = Depends(get_db)):
    new_req = models.AaccupRequirement(
        area_code=req.area_code,
        area_title=req.area_title,
        description=req.description
    )
    db.add(new_req)
    db.commit()
    db.refresh(new_req)
    return new_req

@app.put("/aaccup/requirements/{req_id}", response_model=schemas.AaccupRequirementResponse)
def update_aaccup_requirement(req_id: str, req: schemas.AaccupRequirementCreate, db: Session = Depends(get_db)):
    db_req = db.query(models.AaccupRequirement).filter(models.AaccupRequirement.id == req_id).first()
    if not db_req:
        raise HTTPException(status_code=404, detail="Requirement not found")
    db_req.area_code = req.area_code
    db_req.area_title = req.area_title
    db_req.description = req.description
    db.commit()
    db.refresh(db_req)
    return db_req

@app.delete("/aaccup/requirements/{req_id}")
def delete_aaccup_requirement(req_id: str, db: Session = Depends(get_db)):
    db_req = db.query(models.AaccupRequirement).filter(models.AaccupRequirement.id == req_id).first()
    if not db_req:
        raise HTTPException(status_code=404, detail="Requirement not found")
    db.delete(db_req)
    db.commit()
    return {"message": "Requirement deleted successfully"}


# ─────────────────────────────────────────────────────────────────────────────
# ANALYTICS
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/analytics/recent")
def get_recent_questions():
    """Return recent user queries, aggressively filtering out greetings and small talk."""
    try:
        from vector_store import supabase
        
        # Fetch a large batch so we have enough left after the aggressive filtering
        response = supabase.table("chat_history").select("content").eq("role", "user").order("created_at", desc=True).limit(40).execute()

        valid_recent_questions = []
        
        # The ultimate blacklist of small talk (no punctuation needed here)
        chitchat_blacklist = [
            "hi", "hello", "hey", "good morning", "good afternoon", "good evening", "greetings",
            "how are you", "good morning how are you", "hello how are you", "hi how are you",
            "how are you doing", "what is up", "sup", "hello there", "hi there",
            "thank you", "thanks", "thank you very much", "thanks for the help", 
            "ok", "okay", "yes", "no", "bye", "goodbye"
        ]

        if response.data:
            for row in response.data:
                raw_question = (row.get("content") or "").strip()
                
                # 1. NORMALIZE: Convert to lowercase and completely remove commas, ?, !, and periods
                # "Good morning, how are you?"  -> becomes -> "good morning how are you"
                normalized_q = re.sub(r'[^a-z0-9\s]', '', raw_question.lower()).strip()

                # FILTER 1: Is the normalized text in our small-talk blacklist?
                if normalized_q in chitchat_blacklist:
                    continue
                
                # FILTER 2: Does it start with a greeting but is still very short? 
                # (Catches random things like "hi askpolicy" or "hello chatbot")
                if normalized_q.startswith(("hi ", "hello ", "good morning", "hey ")) and len(normalized_q.split()) <= 4:
                    continue
                    
                # FILTER 3: Is it just one single word? (e.g. "what", "policy", "test")
                if len(normalized_q.split()) < 2:
                    continue
                
                # FILTER 4: Prevent exact duplicates from showing up in the UI
                if raw_question not in valid_recent_questions:
                    valid_recent_questions.append(raw_question)

                # Stop once we have exactly 5 high-quality policy questions
                if len(valid_recent_questions) >= 5:
                    break

        # Fallback if the database has no real questions yet
        if not valid_recent_questions:
            return [
                "What is the grading system?",
                "How do I apply for a scholarship?",
                "What are the requirements for enrollment?"
            ]

        return valid_recent_questions

    except Exception as e:
        print(f"[get_recent_questions] error: {e}")
        return [
            "What is the grading system?",
            "How do I apply for a scholarship?"
        ]


@app.get("/analytics/popular")
def get_popular_topics():
    """Return the top recurring topics from chat history to power the Popular Topics panel."""
    try:
        response = supabase_query_with_retry(lambda: supabase.table("chat_history").select("content").eq("role", "user").execute())

        keyword_map = {
            "enrollment":   ("Enrollment",   "blue"),
            "grade":        ("Grading",      "emerald"),
            "scholarship":  ("Scholarship",  "purple"),
            "attendance":   ("Attendance",   "blue"),
            "graduation":   ("Graduation",   "emerald"),
            "discipline":   ("Discipline",   "purple"),
            "thesis":       ("Thesis",       "blue"),
            "ojt":          ("OJT",          "emerald"),
            "tuition":      ("Tuition",      "purple"),
            "examination":  ("Examination",  "blue"),
            "accreditation":("Accreditation","emerald"),
            "leave":        ("Leave",        "purple"),
        }

        counts = {k: 0 for k in keyword_map}
        if response.data:
            for row in response.data:
                text = (row.get("content") or "").lower()
                for keyword in keyword_map:
                    if keyword in text:
                        counts[keyword] += 1

        # Sort by frequency, return top 6
        sorted_topics = sorted(counts.items(), key=lambda x: x[1], reverse=True)
        top_topics = [
            {"label": keyword_map[k][0], "color": keyword_map[k][1]}
            for k, v in sorted_topics if v > 0
        ][:6]

        # Fallback defaults if no chat history yet
        if not top_topics:
            top_topics = [
                {"label": "Enrollment",  "color": "blue"},
                {"label": "Grading",     "color": "emerald"},
                {"label": "Scholarship", "color": "purple"},
                {"label": "Attendance",  "color": "blue"},
                {"label": "Graduation",  "color": "emerald"},
                {"label": "Discipline",  "color": "purple"},
            ]

        return top_topics

    except Exception as e:
        print(f"[get_popular_topics] error: {e}")
        return [
            {"label": "Enrollment",  "color": "blue"},
            {"label": "Grading",     "color": "emerald"},
            {"label": "Scholarship", "color": "purple"},
        ]


# ─────────────────────────────────────────────────────────────────────────────
# AUDIT TRAIL
# ─────────────────────────────────────────────────────────────────────────────

def format_ph_time(raw_date: str) -> str:
    """Helper to convert UTC ISO strings to Philippine Standard Time (UTC+8)."""
    if not raw_date:
        return "Unknown Date"
    try:
        from datetime import datetime, timezone, timedelta
        # Parse the UTC time
        dt = datetime.fromisoformat(raw_date.replace("Z", "+00:00"))
        
        # Convert to Philippine Time (UTC+8)
        ph_tz = timezone(timedelta(hours=8))
        dt_ph = dt.astimezone(ph_tz)
        
        return dt_ph.strftime("%B %d, %Y - %I:%M %p")
    except Exception:
        return str(raw_date).split("T")[0]


@app.get("/audit/queries")
def get_query_logs():
    try:
        response = supabase_query_with_retry(lambda: supabase.table("chat_history").select("*").eq("role", "user").order("created_at", desc=True).limit(100).execute())

        logs = []
        if response.data:
            for index, item in enumerate(response.data):
                logs.append({
                    "id":        item.get("id", index),
                    "user":      item.get("user_email", "Unknown User"),
                    "role":      item.get("user_role",  item.get("role", "User")),
                    "query":     item.get("content",    ""),
                    "timestamp": format_ph_time(item.get("created_at", "")),
                    "status":    "Answered",
                })

        return logs
    except Exception as e:
        print(f"Audit log error: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch query logs")


@app.post("/audit/access")
def log_document_access(log: AccessLogRequest):
    try:
        supabase.table("document_access_logs").insert({
            "document_name": log.document_name,
            "action_type":   log.action_type,
            "user_email":    log.user_email,
            "user_role":     log.user_role,
        }).execute()
        return {"message": "Access successfully logged"}
    except Exception as e:
        print(f"Failed to log access: {e}")
        return {"message": "Silent failure - do not disrupt user experience"}


@app.get("/audit/access")
def get_access_logs():
    try:
        response = supabase_query_with_retry(lambda: supabase.table("document_access_logs").select("*").order("accessed_at", desc=True).limit(100).execute())

        logs = []
        if response.data:
            for item in response.data:
                logs.append({
                    "id":        item.get("id"),
                    "user":      item.get("user_email",    "Unknown"),
                    "role":      item.get("user_role",     "User"),
                    "document":  item.get("document_name", "Unknown Document"),
                    "action":    item.get("action_type",   "Accessed"),
                    "timestamp": format_ph_time(item.get("accessed_at", "")),
                })
        return logs
    except Exception as e:
        print(f"Access log fetch error: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch access logs")


@app.get("/audit/versions")
def get_version_history():
    try:
        response  = supabase_query_with_retry(lambda: supabase.table("document_sections").select("metadata").execute())
        raw_logs  = []

        if response.data:
            for item in response.data:
                meta     = item.get("metadata", {})
                doc_name = meta.get("name")
                version  = meta.get("version", "1.0")

                if doc_name:
                    raw_date = meta.get("upload_date", "")
                    raw_logs.append({
                        "document":  doc_name,
                        "version":   version,
                        "user":      meta.get("uploaded_by", "System Admin"),
                        "status":    meta.get("status", "Active"),
                        "timestamp": format_ph_time(raw_date),
                        "raw_date":  raw_date,
                    })

        raw_logs.sort(key=lambda x: x["raw_date"], reverse=True)

        unique_logs = []
        seen        = set()
        for idx, log in enumerate(raw_logs):
            identifier = f"{log['document']}_v{log['version']}"
            if identifier not in seen:
                seen.add(identifier)
                log["id"] = f"ver_{idx}"
                unique_logs.append(log)

        return unique_logs[:100]

    except Exception as e:
        print(f"Version log fetch error: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch version history")


@app.get("/audit/system")
def get_system_events():
    try:
        response = supabase_query_with_retry(lambda: supabase.table("system_events_logs").select("*").order("event_date", desc=True).limit(100).execute())

        logs = []
        if response.data:
            for item in response.data:
                logs.append({
                    "id":          item.get("id"),
                    "user":        item.get("user_email",  "Unknown"),
                    "type":        item.get("event_type",  "System Event"),
                    "description": item.get("description", ""),
                    "timestamp":   format_ph_time(item.get("event_date", "")),
                })
        return logs
    except Exception as e:
        print(f"System log fetch error: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch system events")

# ─────────────────────────────────────────────────────────────────────────────
# GRADE EVALUATION
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/evaluate-grades")
async def evaluate_grades(file: UploadFile = File(...)):
    try:
        content    = await file.read()
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        raw_text   = ""
        for page in pdf_reader.pages:
            raw_text += page.extract_text() + "\n"

        system_prompt = """
        You are a meticulous Academic Data Extractor for Cebu Technological University (CTU).
        Your job is to parse scrambled PDF text and extract subjects, units, and grades into a clean JSON structure.

        CRITICAL PARSING RULES FOR PDFs:
        1. PyPDF2 flattens tables. A subject row might look like: "CS46 FUNCTIONAL ENGLISH 3.00 01:00PM Wed CAS 3 1.6"
        2. UNITS are typically integers or decimals like 2.00, 3.00, or 5.00.
        3. GRADES are strictly between 1.0 and 5.0.
        4. DO NOT INVENT GRADES. If a subject does not have a clear grade (1.0 to 5.0) explicitly associated with it on its line, mark the grade as 0 and 'has_missing_grades' as true.
        5. If there are multiple grades on a single line, ALWAYS extract the LAST valid numerical grade on that line.

        You MUST respond with a pure JSON object in this EXACT format.
        {
          "semesters": [
            {
              "semester_name": "1st Semester SY 2023-2024",
              "subjects_scratchpad": [
                {"subject": "CS46 FUNCTIONAL ENGLISH", "units": 3.0, "grade": 1.6, "weighted_score": 4.8},
                {"subject": "CS47 ART APP", "units": 3.0, "grade": 0, "weighted_score": 0}
              ],
              "has_missing_grades": true
            }
          ],
          "summary": "2-3 sentences summarizing performance. Be direct using 'You'.",
          "recommendations": ["Recommendation 1", "Recommendation 2", "Recommendation 3"]
        }
        """

        response    = groq_client.chat.completions.create(
            model="qwen2.5",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": f"Here is the raw text from the grade slip:\n\n{raw_text}"}
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )

        result_json = response.choices[0].message.content
        return json.loads(result_json)

    except Exception as e:
        print(f"Grade Evaluation Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to evaluate grades.")


# ─────────────────────────────────────────────────────────────────────────────
# NOTIFICATION SYSTEM ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/notifications", response_model=List[NotificationOut])
def get_notifications(email: str, filter: str = "all"):
    if not email:
        raise HTTPException(status_code=400, detail="email query param is required")

    try:
        def _query():
            q = (
                supabase.table("notifications")
                .select("*")
                .eq("user_email", email)
                .order("created_at", desc=True)
                .limit(100)
            )
            if filter == "unread":
                q = q.eq("is_read", False)
            return q.execute()

        response = supabase_query_with_retry(_query)
        return response.data if response.data else []

    except Exception as e:
        print(f"[notifications] GET error: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch notifications")


@app.get("/notifications/count")
def get_notification_count(email: str):
    if not email:
        raise HTTPException(status_code=400, detail="email query param is required")

    try:
        response = (
            supabase.table("notifications")
            .select("id", count="exact")
            .eq("user_email", email)
            .eq("is_read", False)
            .execute()
        )
        return {"unread_count": response.count if response.count is not None else 0}

    except Exception as e:
        print(f"[notifications] COUNT error: {e}")
        raise HTTPException(status_code=500, detail="Failed to count notifications")


@app.patch("/notifications/mark-all-read")
def mark_all_notifications_read(payload: dict = Body(...), db: Session = Depends(get_db)):
    email = payload.get("email")
    if not email:
        raise HTTPException(status_code=400, detail="email is required in request body")

    try:
        response = (
            supabase.table("notifications")
            .update({"is_read": True})
            .eq("user_email", email)
            .eq("is_read", False)
            .execute()
        )
        updated = len(response.data) if response.data else 0

        # Increment read_count for any broadcast announcements present in read notifications
        if response.data:
            for item in response.data:
                ntitle = item.get("title", "")
                if "📢" in ntitle or "Announcement" in ntitle:
                    clean = ntitle.replace("📢", "").strip()
                    matched = db.query(models.Announcement).filter(models.Announcement.title.ilike(f"%{clean}%")).first()
                    if matched:
                        matched.read_count = (matched.read_count or 0) + 1
            db.commit()

        return {"message": f"Marked {updated} notification(s) as read", "updated": updated}

    except Exception as e:
        print(f"[notifications] MARK-ALL-READ error: {e}")
        raise HTTPException(status_code=500, detail="Failed to mark notifications")


@app.patch("/notifications/{notification_id}/read")
def mark_notification_read(notification_id: int, db: Session = Depends(get_db)):
    try:
        response = (
            supabase.table("notifications")
            .update({"is_read": True})
            .eq("id", notification_id)
            .execute()
        )
        if not response.data:
            raise HTTPException(status_code=404, detail="Notification not found")

        # Sync read status to matching announcement if applicable
        notif_item = response.data[0]
        ntitle = notif_item.get("title", "")
        if "📢" in ntitle or "Announcement" in ntitle:
            clean = ntitle.replace("📢", "").strip()
            matched = db.query(models.Announcement).filter(models.Announcement.title.ilike(f"%{clean}%")).first()
            if matched:
                matched.read_count = (matched.read_count or 0) + 1
                db.commit()

        return {"message": "Notification marked as read"}

    except HTTPException:
        raise
    except Exception as e:
        print(f"[notifications] MARK-READ error: {e}")
        raise HTTPException(status_code=500, detail="Failed to update notification")


@app.delete("/notifications/delete-read")
def delete_read_notifications(payload: dict = Body(...)):
    email = payload.get("email")
    if not email:
        raise HTTPException(status_code=400, detail="email is required in request body")

    try:
        response = (
            supabase.table("notifications")
            .delete()
            .eq("user_email", email)
            .eq("is_read", True)
            .execute()
        )
        deleted = len(response.data) if response.data else 0
        return {"message": f"Deleted {deleted} read notification(s)", "deleted": deleted}

    except Exception as e:
        print(f"[notifications] DELETE-READ error: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete read notifications")


@app.post("/notifications", response_model=NotificationOut, status_code=201)
def create_notification_endpoint(notification: NotificationCreate):
    if notification.type not in VALID_NOTIF_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"type must be one of: {', '.join(VALID_NOTIF_TYPES)}"
        )

    try:
        response = supabase.table("notifications").insert({
            "user_email": notification.user_email,
            "type":       notification.type,
            "title":      notification.title,
            "message":    notification.message,
            "is_read":    False,
        }).execute()
        return response.data[0]

    except Exception as e:
        print(f"[notifications] CREATE error: {e}")
        raise HTTPException(status_code=500, detail="Failed to create notification")
    
@app.post("/users/change-password")
def change_password(req: ChangePasswordRequest, db: Session = Depends(get_db)):
    # 1. Find the user
    user = db.query(models.User).filter(models.User.email == req.email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # 2. Verify their current password (Security Check)
    if not utils.verify_password(req.current_password, user.hashed_password):
        raise HTTPException(status_code=400, detail="Incorrect current password")

    # 3. Hash and save the new password
    user.hashed_password = utils.hash_password(req.new_password)
    db.commit()

    # --- SILENT AUDIT LOG ---
    try:
        from vector_store import supabase
        supabase.table("system_events_logs").insert({
            "user_email": req.email,
            "event_type": "Security Update",
            "description": "User successfully changed their password"
        }).execute()
    except Exception as e:
        print(f"Failed to log password change: {e}")
    # ------------------------
    
    return {"message": "Password successfully updated!"}

@app.put("/users/profile")
def update_profile(req: UpdateProfileRequest, db: Session = Depends(get_db)):
    # 1. Fetch the core user
    user = db.query(models.User).filter(models.User.email == req.email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # 2. Handle Email Change Logic
    if req.new_email != req.email:
        existing_email = db.query(models.User).filter(models.User.email == req.new_email).first()
        if existing_email:
            raise HTTPException(status_code=400, detail="This email is already in use by another account.")
        user.email = req.new_email
    
    # 3. Update core user details
    user.full_name = req.full_name
    
    # 4. Relational Table Routing
    if user.role == "STUDENT":
        if user.student_profile:
            user.student_profile.course = req.program
        else:
            new_profile = models.StudentProfile(user_id=user.id, course=req.program)
            db.add(new_profile)
    else:
        user.department = req.program
        
    db.commit()

    # --- SILENT AUDIT LOG ---
    try:
        from vector_store import supabase
        supabase.table("system_events_logs").insert({
            "user_email": user.email, # Use the confirmed email
            "event_type": "Profile Update",
            "description": f"Updated profile details (Name: {req.full_name}, Program: {req.program})"
        }).execute()
    except Exception as e:
        print(f"Failed to log profile update: {e}")
    # ------------------------
    
    return {
        "message": "Profile updated successfully!", 
        "full_name": user.full_name, 
        "program": req.program,
        "email": user.email  
    }

# --- NEW: ANNOUNCEMENT ROUTES ---
@app.post("/announcements", response_model=schemas.AnnouncementResponse)
def create_announcement(announcement: schemas.AnnouncementCreate, db: Session = Depends(get_db)):
    sent_dt = datetime.utcnow()
    
    # If it's scheduled, parse the HTML datetime string
    if announcement.schedule_date:
        try:
            sent_dt = datetime.fromisoformat(announcement.schedule_date.replace("Z", "+00:00"))
        except ValueError:
            pass # Fallback to current time if parsing fails
            
    db_announcement = models.Announcement(
        title=announcement.title,
        content=announcement.content,
        recipients=announcement.recipients,
        sent_date=sent_dt,
        sent_by=announcement.sent_by,
        status=announcement.status,
        total_recipients=announcement.total_recipients
    )
    
    db.add(db_announcement)
    db.commit()
    db.refresh(db_announcement)
    
    # --- SILENT AUDIT LOG ---
    try:
        from vector_store import supabase
        supabase.table("system_events_logs").insert({
            "user_email": announcement.sent_by,
            "event_type": "Broadcast Sent",
            "description": f"Broadcasted: {announcement.title} to {announcement.recipients}"
        }).execute()
    except Exception as e:
        print(f"Failed to log announcement: {e}")

    return db_announcement

@app.get("/announcements", response_model=List[schemas.AnnouncementResponse])
def get_announcements(db: Session = Depends(get_db)):
    # Fetch all announcements, newest first
    return db.query(models.Announcement).order_by(models.Announcement.sent_date.desc()).all()

@app.get("/users/counts")
def get_user_counts(db: Session = Depends(get_db)):
    # Fetch real-time counts from the database, ignoring disabled accounts
    students = db.query(models.User).filter(models.User.role == "STUDENT", models.User.status == "Active").count()
    faculty = db.query(models.User).filter(models.User.role == "FACULTY", models.User.status == "Active").count()
    admins = db.query(models.User).filter(models.User.role == "ADMIN", models.User.status == "Active").count()
    
    total = students + faculty + admins
    return {
        "all": total,
        "students": students,
        "faculty": faculty
    }

@app.put("/announcements/{announcement_id}", response_model=schemas.AnnouncementResponse)
def update_announcement(announcement_id: str, req: schemas.AnnouncementUpdate, db: Session = Depends(get_db)):
    announcement = db.query(models.Announcement).filter(models.Announcement.id == announcement_id).first()
    if not announcement:
        raise HTTPException(status_code=404, detail="Announcement not found")
    
    announcement.title = req.title
    announcement.content = req.content
    announcement.recipients = req.recipients
    announcement.status = req.status
    announcement.total_recipients = req.total_recipients
    
    if req.schedule_date:
        try:
            announcement.sent_date = datetime.fromisoformat(req.schedule_date.replace("Z", "+00:00"))
        except ValueError:
            pass 
    elif req.status == "Sent":
        announcement.sent_date = datetime.utcnow() # Update timestamp if sending right now

    db.commit()
    db.refresh(announcement)
    return announcement

@app.delete("/announcements/{announcement_id}")
def delete_announcement(announcement_id: str, db: Session = Depends(get_db)):
    announcement = db.query(models.Announcement).filter(models.Announcement.id == announcement_id).first()
    if not announcement:
        raise HTTPException(status_code=404, detail="Announcement not found")
    
    # Security check: Prevent deleting Sent announcements via API
    if announcement.status == "Sent":
        raise HTTPException(status_code=400, detail="Cannot delete an announcement that has already been sent.")
        
    db.delete(announcement)
    db.commit()
    return {"message": "Announcement deleted successfully."}

@app.post("/announcements/{announcement_id}/read")
def mark_announcement_read(announcement_id: str, db: Session = Depends(get_db)):
    announcement = db.query(models.Announcement).filter(models.Announcement.id == announcement_id).first()
    if not announcement:
        raise HTTPException(status_code=404, detail="Announcement not found")
    
    announcement.read_count = (announcement.read_count or 0) + 1
    db.commit()
    db.refresh(announcement)
    return {
        "id": str(announcement.id),
        "read_count": announcement.read_count,
        "total_recipients": announcement.total_recipients
    }

# --- SETTINGS ROUTES ---
@app.get("/settings", response_model=SettingsSchema)
def get_system_settings(db: Session = Depends(get_db)):
    settings = db.query(models.SystemSettings).filter(models.SystemSettings.id == 1).first()
    
    # If settings don't exist yet, create the default row
    if not settings:
        settings = models.SystemSettings(id=1)
        db.add(settings)
        db.commit()
        db.refresh(settings)
        
    return settings

@app.put("/settings")
def update_system_settings(req: SettingsSchema, db: Session = Depends(get_db)):
    settings = db.query(models.SystemSettings).filter(models.SystemSettings.id == 1).first()
    if not settings:
        settings = models.SystemSettings(id=1)
        db.add(settings)
    
    # Update all fields
    settings.platform_name = req.platform_name
    settings.campus = req.campus
    settings.admin_email = req.admin_email
    settings.jwt_expiration = req.jwt_expiration
    settings.otp_expiration = req.otp_expiration
    settings.ai_model = req.ai_model
    settings.ai_temperature = req.ai_temperature
    settings.ai_system_prompt = req.ai_system_prompt
    settings.rag_max_chunks = req.rag_max_chunks
    
    db.commit()
    
    # --- SILENT AUDIT LOG ---
    try:
        from vector_store import supabase
        supabase.table("system_events_logs").insert({
            "user_email": "System Admin",
            "event_type": "System Config Update",
            "description": "Administrator modified core system and AI settings."
        }).execute()
    except Exception as e:
        pass
    
    return {"message": "Settings successfully updated!"}

# ─────────────────────────────────────────────────────────────────────────────
# CHED MONITORING MODULE
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/ched/requirements", response_model=schemas.ChedRequirementResponse)
def create_ched_requirement(req: schemas.ChedRequirementCreate, db: Session = Depends(get_db)):
    """Admin endpoint to add a new blank requirement to the checklist."""
    new_req = models.ChedRequirement(
        program=req.program,
        cmo_name=req.cmo_name,
        description=req.description,
        status="Not Compliant"
    )
    db.add(new_req)
    db.commit()
    db.refresh(new_req)
    
    # --- SILENT AUDIT LOG ---
    try:
        supabase.table("system_events_logs").insert({
            "user_email": "System Admin",
            "event_type": "CHED Setup",
            "description": f"Added new CHED requirement for {req.program}"
        }).execute()
    except Exception:
        pass
        
    return new_req

@app.put("/ched/requirements/{req_id}", response_model=schemas.ChedRequirementResponse)
def update_ched_requirement(req_id: str, req: schemas.ChedRequirementCreate, db: Session = Depends(get_db)):
    """Admin endpoint to edit an existing requirement."""
    requirement = db.query(models.ChedRequirement).filter(models.ChedRequirement.id == req_id).first()
    if not requirement:
        raise HTTPException(status_code=404, detail="Requirement not found")
        
    requirement.cmo_name = req.cmo_name
    requirement.description = req.description
    db.commit()
    db.refresh(requirement)
    return requirement

@app.delete("/ched/requirements/{req_id}")
def delete_ched_requirement(req_id: str, db: Session = Depends(get_db)):
    """Admin endpoint to delete a requirement."""
    requirement = db.query(models.ChedRequirement).filter(models.ChedRequirement.id == req_id).first()
    if not requirement:
        raise HTTPException(status_code=404, detail="Requirement not found")
        
    db.delete(requirement)
    db.commit()
    return {"message": "Requirement deleted successfully."}

@app.get("/ched/requirements/{program}", response_model=List[schemas.ChedRequirementResponse])
def get_ched_requirements(program: str, db: Session = Depends(get_db)):
    """Fetches all requirements for a specific program (e.g., BSIT), including nested attached files."""
    reqs = db.query(models.ChedRequirement).filter(models.ChedRequirement.program == program).all()
    return reqs

@app.post("/ched/upload-evidence")
async def upload_ched_evidence(
    file: UploadFile = File(...),
    requirement_id: str = Form(...),
    document_name: str = Form(...),
    uploaded_by: str = Form(...),
    program: str = Form(...), # NEW: Needed for vector metadata
    db: Session = Depends(get_db)
):
    """Uploads PDF evidence to Supabase Storage, links it to CHED, and embeds it for RAG."""
    import PyPDF2
    
    # 1. Verify the requirement exists
    requirement = db.query(models.ChedRequirement).filter(models.ChedRequirement.id == requirement_id).first()
    if not requirement:
        raise HTTPException(status_code=404, detail="CHED Requirement not found")

    # 2. Extract Text & Upload File
    contents = await file.read()
    extracted_text = ""
    filename_lower = file.filename.lower()

    if filename_lower.endswith(".pdf"):
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(contents))
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text: extracted_text += text + "\n"
    elif filename_lower.endswith(".docx"):
        import docx
        doc = docx.Document(io.BytesIO(contents))
        extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    safe_filename = file.filename.replace(" ", "_")
    unique_filename = f"ched_{int(time.time())}_{safe_filename}"
    
    try:
        supabase.storage.from_("documents").upload(
            file=contents, path=unique_filename, file_options={"content-type": file.content_type}
        )
        public_url = supabase.storage.from_("documents").get_public_url(unique_filename)
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to upload file to storage.")

    # 3. Save Record in PostgreSQL Database
    new_evidence = models.ChedEvidence(
        requirement_id=requirement.id,
        document_name=document_name,
        file_url=public_url,
        uploaded_by=uploaded_by
    )
    db.add(new_evidence)
    
    # 4. Automatically change the requirement status to "Pending"
    requirement.status = "Pending"
    db.commit()

    # 5. NEW: Embed into Vector Database so the AI can read it!
    if extracted_text.strip():
        metadata = {
            "name": document_name,
            "category": "CHED Evidence",
            "office": "Quality Assurance",
            "version": "1.0",
            "status": "Pending",
            "program": program,
            "requirement_target": requirement.description,
            "uploaded_by": uploaded_by,
            "upload_date": datetime.now().isoformat(),
            "file_url": public_url,
            "is_ched": True,
            "req_id": str(requirement.id) # Link back to the SQL row
        }
        vector_store.add_to_vector_db(extracted_text, metadata)

    return {"message": "Evidence uploaded successfully. Status set to Pending Review."}

@app.put("/ched/requirements/{requirement_id}/status")
def update_ched_status(requirement_id: str, status: str = Body(..., embed=True), db: Session = Depends(get_db)):
    """Admin endpoint to Accept or Revoke. If Revoked, deletes the attached evidence."""
    req = db.query(models.ChedRequirement).filter(models.ChedRequirement.id == requirement_id).first()
    if not req:
        raise HTTPException(status_code=404, detail="Requirement not found")
        
    req.status = status 
    
    # If the admin Revokes or Rejects, we must delete the evidence file record
    if status == "Not Compliant":
        # Delete from SQL
        db.query(models.ChedEvidence).filter(models.ChedEvidence.requirement_id == requirement_id).delete()
        
        # Archive from Vector DB so AI stops reading it
        try:
            from vector_store import supabase
            chunks_res = supabase.table("document_sections").select("id, metadata").eq("metadata->>req_id", requirement_id).execute()
            if chunks_res.data:
                for chunk in chunks_res.data:
                    chunk_meta = chunk['metadata']
                    chunk_meta['status'] = "Archived"
                    supabase.table("document_sections").update({"metadata": chunk_meta}).eq("id", chunk['id']).execute()
        except Exception:
            pass
            
    db.commit()
    return {"message": f"Requirement successfully marked as {status}"}

@app.put("/ched/evidence/{evidence_id}/status")
def update_ched_evidence_status(evidence_id: str, status: str = Body(..., embed=True), db: Session = Depends(get_db)):
    """Admin endpoint to update status of a CHED evidence file."""
    evidence = db.query(models.ChedEvidence).filter(models.ChedEvidence.id == evidence_id).first()
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
    evidence.status = status
    db.commit()
    return {"message": f"Evidence status updated to {status}"}

# --- NEW: DELETE CHED EVIDENCE ---
@app.delete("/ched/evidence/{evidence_id}")
def delete_ched_evidence(evidence_id: str, db: Session = Depends(get_db)):
    """Deletes specific CHED evidence and archives its vector chunks."""
    evidence = db.query(models.ChedEvidence).filter(models.ChedEvidence.id == evidence_id).first()
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
        
    req_id = evidence.requirement_id
    doc_name = evidence.document_name
    
    # Delete from SQL
    db.delete(evidence)
    
    # Auto-revert requirement status to Not Compliant if empty
    req = db.query(models.ChedRequirement).filter(models.ChedRequirement.id == req_id).first()
    if req and len(req.evidences) == 0:
        req.status = "Not Compliant"
        
    db.commit()
    
    # Archive from Vector DB
    try:
        from vector_store import supabase
        chunks_res = supabase.table("document_sections").select("id, metadata").eq("metadata->>req_id", str(req_id)).eq("metadata->>name", doc_name).execute()
        if chunks_res.data:
            for chunk in chunks_res.data:
                chunk_meta = chunk['metadata']
                chunk_meta['status'] = "Archived"
                supabase.table("document_sections").update({"metadata": chunk_meta}).eq("id", chunk['id']).execute()
    except Exception:
        pass
        
    return {"message": "Evidence deleted successfully."}


# ─────────────────────────────────────────────────────────────────────────────
# PAPER TRAIL (RECEIVING & RELEASING HISTORY) API ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────

def _generate_tracking_number(db: Session) -> str:
    """Generates a unique tracking number format: PT-YYYY-XXXX."""
    import random
    year = datetime.now().year
    while True:
        num = random.randint(1000, 9999)
        tracking_no = f"PT-{year}-{num}"
        existing = db.query(models.PaperTrailRecord).filter(models.PaperTrailRecord.tracking_number == tracking_no).first()
        if not existing:
            return tracking_no


@app.post("/paper-trail", response_model=schemas.PaperTrailRecordResponse, status_code=status.HTTP_201_CREATED)
def create_paper_trail_record(
    payload: schemas.PaperTrailCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_faculty_or_admin)
):
    """Creates a new document paper trail record and logs initial release."""
    tracking_no = _generate_tracking_number(db)
    
    sender_name = current_user.full_name or payload.sender_name
    sender_email = current_user.email
    sender_role = current_user.role.upper()

    new_record = models.PaperTrailRecord(
        tracking_number=tracking_no,
        title=payload.title,
        document_type=payload.document_type,
        office=payload.office,
        sender_name=sender_name,
        sender_email=sender_email,
        sender_role=sender_role,
        recipient_name=payload.recipient_name,
        recipient_email=payload.recipient_email,
        recipient_role=payload.recipient_role.upper() if payload.recipient_role else None,
        origin_office=payload.origin_office or getattr(current_user, 'administrative_office', None) or getattr(current_user, 'department', None) or payload.office,
        origin_person=payload.origin_person or sender_name,
        current_location=payload.current_location or payload.office,
        transaction_type=payload.transaction_type or "Submission",
        status="Pending Receiving",
        remarks=payload.remarks,
        file_url=payload.file_url
    )
    db.add(new_record)
    db.commit()
    db.refresh(new_record)

    # Initial Log Entry
    initial_log = models.PaperTrailLog(
        record_id=new_record.id,
        action="Document Released / Submitted",
        status="Pending Receiving",
        actor_name=sender_name,
        actor_email=sender_email,
        actor_role=sender_role,
        notes=payload.remarks or f"Document '{payload.title}' released to {payload.office}."
    )
    db.add(initial_log)
    db.commit()
    db.refresh(new_record)

    # Notifications
    try:
        # Office Broadcasting to target office
        _notify_office(
            db=db,
            office_name=payload.office,
            n_type="info",
            title=f"New Document Received: {tracking_no}",
            message=f"{sender_name} released document '{payload.title}' to your office ({payload.office})."
        )
        if payload.recipient_email and payload.recipient_email != sender_email:
            _send_notification(
                user_email=payload.recipient_email,
                n_type="info",
                title=f"New Document Received: {tracking_no}",
                message=f"{sender_name} released document '{payload.title}' to your office ({payload.office})."
            )
        _notify_all_admins(
            db=db,
            n_type="info",
            title=f"Paper Trail Created: {tracking_no}",
            message=f"Document '{payload.title}' ({payload.document_type}) released by {sender_name} to {payload.office}."
        )
    except Exception as exc:
        print(f"[paper_trail] notification warning: {exc}")

    return new_record


@app.post("/paper-trail/request", response_model=schemas.PaperTrailRecordResponse, status_code=status.HTTP_201_CREATED)
def create_top_down_document_request(
    payload: schemas.PaperTrailRequestCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_faculty_or_admin)
):
    """Faculty or Admin/Director creates a document request assigned to an office or person (e.g. Memo, Resolution)."""
    tracking_no = _generate_tracking_number(db)
    requester_name = current_user.full_name or "Faculty / User"
    requester_email = current_user.email
    requester_role = current_user.role.upper()
    requester_office = getattr(current_user, 'administrative_office', None) or getattr(current_user, 'department', None) or payload.office

    new_record = models.PaperTrailRecord(
        tracking_number=tracking_no,
        title=payload.title,
        document_type=payload.document_type,
        office=payload.office,
        sender_name=requester_name,
        sender_email=requester_email,
        sender_role=requester_role,
        recipient_name=payload.target_person_name,
        recipient_email=payload.target_person_email,
        recipient_role="FACULTY" if requester_role == "ADMIN" else "ADMIN",
        origin_office=requester_office,
        origin_person=requester_name,
        current_location=payload.office or f"Desk of {payload.target_person_name}",
        transaction_type="Request",
        status="Pending Request",
        remarks=payload.instructions or f"Request for {payload.document_type}: {payload.title}"
    )
    db.add(new_record)
    db.commit()
    db.refresh(new_record)

    initial_log = models.PaperTrailLog(
        record_id=new_record.id,
        action=f"Document Request Issued ({payload.document_type})",
        status="Pending Request",
        actor_name=requester_name,
        actor_email=requester_email,
        actor_role=requester_role,
        notes=payload.instructions or f"{requester_name} requested '{payload.title}' ({payload.document_type}) from {payload.office}."
    )
    db.add(initial_log)
    db.commit()

    try:
        # Notify target person if email provided
        if payload.target_person_email:
            _send_notification(
                user_email=payload.target_person_email,
                n_type="warning",
                title=f"Pending Document Request [{tracking_no}]",
                message=f"{requester_name} has requested document '{payload.title}' ({payload.document_type}). Please upload and fulfill."
            )
        # Office Broadcasting to target office staff
        _notify_office(
            db=db,
            office_name=payload.office,
            n_type="warning",
            title=f"Document Request Issued [{tracking_no}]",
            message=f"{requester_name} requested document '{payload.title}' ({payload.document_type}) from {payload.office}."
        )
    except Exception as exc:
        print(f"[create_top_down_document_request] notification warning: {exc}")

    return new_record


@app.put("/paper-trail/{record_id}/fulfill", response_model=schemas.PaperTrailRecordResponse)
def fulfill_document_request(
    record_id: str,
    payload: schemas.PaperTrailFulfillRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_faculty_or_admin)
):
    """Faculty fulfills a top-down document request, attaching the file and routing it back to origin Director/Admin."""
    record = db.query(models.PaperTrailRecord).filter(models.PaperTrailRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Document record not found.")

    user_name = current_user.full_name or "Faculty Member"
    user_email = current_user.email

    record.file_url = payload.file_url
    record.status = "Submitted / Fulfilled"
    record.current_location = record.origin_office or record.sender_name or "Director / Admin Desk"
    record.updated_at = datetime.utcnow()

    new_log = models.PaperTrailLog(
        record_id=record.id,
        action="Request Fulfilled & Routed to Origin",
        status="Submitted / Fulfilled",
        actor_name=user_name,
        actor_email=user_email,
        actor_role=current_user.role.upper(),
        notes=payload.remarks or f"Fulfilled request for '{record.title}' and routed back to {record.current_location}."
    )
    db.add(new_log)
    db.commit()
    db.refresh(record)

    try:
        # Requester / Originator Tracking notification
        if record.sender_email:
            _send_notification(
                user_email=record.sender_email,
                n_type="success",
                title=f"Document Request Fulfilled [{record.tracking_number}]",
                message=f"{user_name} has fulfilled the request for '{record.title}' and routed it back to your desk."
            )
        # Office Broadcasting to origin office
        if record.origin_office:
            _notify_office(
                db=db,
                office_name=record.origin_office,
                n_type="success",
                title=f"Document Request Fulfilled [{record.tracking_number}]",
                message=f"{user_name} has fulfilled the request for '{record.title}'."
            )
    except Exception as exc:
        print(f"[fulfill_document_request] notification warning: {exc}")

    return record


@app.get("/paper-trail", response_model=List[schemas.PaperTrailRecordResponse])
def get_paper_trail_records(
    role: Optional[str] = None,
    email: Optional[str] = None,
    office: Optional[str] = None,
    status_filter: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_faculty_or_admin)
):
    """Fetches paper trail records enforcing strict role-based data privacy."""
    query = db.query(models.PaperTrailRecord)
    
    user_role = current_user.role.upper()
    user_email = current_user.email
    user_admin_office = getattr(current_user, 'administrative_office', None) or ""
    user_dept = getattr(current_user, 'department', None) or ""
    
    # Strict Data Privacy: If FACULTY (not Admin), ONLY return records where:
    # 1. sender_email matches user's email
    # 2. recipient_email matches user's email
    # 3. office, current_location, or origin_office matches user's administrative_office or department
    if user_role == "FACULTY":
        conditions = [
            models.PaperTrailRecord.sender_email == user_email,
            models.PaperTrailRecord.recipient_email == user_email,
        ]
        if user_admin_office:
            conditions.extend([
                models.PaperTrailRecord.office == user_admin_office,
                models.PaperTrailRecord.current_location == user_admin_office,
                models.PaperTrailRecord.origin_office == user_admin_office,
            ])
        if user_dept:
            conditions.extend([
                models.PaperTrailRecord.office == user_dept,
                models.PaperTrailRecord.current_location == user_dept,
                models.PaperTrailRecord.origin_office == user_dept,
            ])
        query = query.filter(or_(*conditions))
    
    if office and office != "all":
        query = query.filter(models.PaperTrailRecord.office == office)
        
    if status_filter and status_filter != "all":
        query = query.filter(models.PaperTrailRecord.status == status_filter)
        
    return query.order_by(models.PaperTrailRecord.updated_at.desc()).all()


@app.get("/paper-trail/{record_id}", response_model=schemas.PaperTrailRecordResponse)
def get_paper_trail_detail(
    record_id: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_faculty_or_admin)
):
    """Fetches single paper trail record with full movement history."""
    record = db.query(models.PaperTrailRecord).filter(models.PaperTrailRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Paper trail record not found.")
    return record


@app.put("/paper-trail/{record_id}/status", response_model=schemas.PaperTrailRecordResponse)
def update_paper_trail_status(
    record_id: str,
    payload: schemas.PaperTrailStatusUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_faculty_or_admin)
):
    """Updates document routing action & location (Receive, Forward/Route, Return for Revision, Approve & Close)."""
    record = db.query(models.PaperTrailRecord).filter(models.PaperTrailRecord.id == record_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="Document record not found.")

    old_status = record.status
    action_type = payload.action_type or "Acknowledge"
    new_status = payload.status

    if action_type == "Acknowledge" or new_status == "Received":
        new_status = "Received"
        action_text = f"Document Received & Acknowledged at {record.current_location or record.office}"
    elif action_type == "Forward":
        new_status = "Forwarded"
        target_loc = payload.target_office or payload.target_person or "Target Office"
        record.current_location = target_loc
        action_text = f"Forwarded / Routed to {target_loc}"
    elif action_type == "Return" or new_status == "Needs Revision":
        new_status = "Needs Revision"
        target_loc = record.origin_office or record.sender_name or "Origin Office"
        record.current_location = target_loc
        action_text = f"Returned for Revision to {target_loc}"
    elif action_type == "Approve" or new_status == "Approved":
        new_status = "Approved"
        action_text = f"Approved & Closed at {record.current_location or record.office}"
    else:
        action_text = f"Status updated to {new_status}"

    record.status = new_status
    record.updated_at = datetime.utcnow()

    # Append movement log
    new_log = models.PaperTrailLog(
        record_id=record.id,
        action=action_text,
        status=new_status,
        actor_name=payload.actor_name,
        actor_email=payload.actor_email,
        actor_role=payload.actor_role.upper(),
        notes=payload.notes or f"Action: {action_text} by {payload.actor_name}."
    )
    db.add(new_log)
    db.commit()
    db.refresh(record)

    # Notifications (Requester Tracking & Office Broadcasting)
    try:
        # 1. Requester / Originator Tracking: Always notify original person who started trail
        if record.sender_email:
            _send_notification(
                user_email=record.sender_email,
                n_type="info" if new_status != "Needs Revision" else "warning",
                title=f"DTS Movement [{record.tracking_number}]",
                message=f"Document '{record.title}' action: {action_text} by {payload.actor_name}."
            )

        # 2. Office Broadcasting to target/current office
        if action_type == "Forward" and payload.target_office:
            _notify_office(
                db=db,
                office_name=payload.target_office,
                n_type="info",
                title=f"Document Forwarded to Office [{record.tracking_number}]",
                message=f"Document '{record.title}' has been forwarded to {payload.target_office} by {payload.actor_name}."
            )
        elif action_type == "Return" and record.origin_office:
            _notify_office(
                db=db,
                office_name=record.origin_office,
                n_type="warning",
                title=f"Document Returned for Revision [{record.tracking_number}]",
                message=f"Document '{record.title}' was returned to {record.origin_office} by {payload.actor_name}."
            )
        elif action_type == "Approve":
            _notify_office(
                db=db,
                office_name=record.current_location or record.office,
                n_type="success",
                title=f"Document Approved & Closed [{record.tracking_number}]",
                message=f"Document '{record.title}' was verified, approved, and closed by {payload.actor_name}."
            )
    except Exception as exc:
        print(f"[update_paper_trail_status] notification error: {exc}")

    return record


@app.post("/paper-trail/upload")
async def upload_paper_trail_attachment(file: UploadFile = File(...)):
    """Uploads an optional file attachment for a paper trail record."""
    try:
        contents = await file.read()
        safe_filename = file.filename.replace(" ", "_")
        unique_filename = f"papertrail/{int(time.time())}_{safe_filename}"

        supabase.storage.from_("documents").upload(
            file=contents,
            path=unique_filename,
            file_options={"content-type": file.content_type or "application/pdf"}
        )
        public_url = supabase.storage.from_("documents").get_public_url(unique_filename)
        return {"file_url": public_url, "filename": file.filename}
    except Exception as exc:
        print(f"[upload_paper_trail_attachment] error: {exc}")
        raise HTTPException(status_code=500, detail="Failed to upload attachment.")


# ─────────────────────────────────────────────────────────────────────────────
# ISO 9001:2015 QUALITY MANAGEMENT SYSTEM (QMS) & IQA ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/iso/cycles", response_model=List[str])
def get_iso_cycles(db: Session = Depends(get_db)):
    """Retrieves all distinct audit cycle years present in the system, sorted."""
    cycles = db.query(models.ISORequirement.cycle_year).distinct().all()
    cycle_list = [c[0] for c in cycles if c[0]]
    defaults = ["2026 Recertification", "2025 Surveillance", "2024 Initial Audit"]
    for d in defaults:
        if d not in cycle_list:
            cycle_list.append(d)
    return sorted(cycle_list, reverse=True)


@app.post("/iso/cycles/init", response_model=List[schemas.ISORequirementResponse])
def initialize_iso_cycle(
    payload: schemas.ISOCycleCreate,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Initializes a new ISO Audit Cycle (e.g. '2026 Surveillance')."""
    target_prog = "GLOBAL"
    new_cycle = payload.cycle_year.strip()
    if not new_cycle:
        raise HTTPException(status_code=400, detail="Cycle year cannot be empty.")

    existing = db.query(models.ISORequirement).filter(
        models.ISORequirement.program == target_prog,
        models.ISORequirement.cycle_year == new_cycle
    ).all()

    return existing


@app.get("/iso/requirements/{program}", response_model=List[schemas.ISORequirementResponse])
def get_iso_requirements(program: str, cycle_year: Optional[str] = "2025 Surveillance", db: Session = Depends(get_db)):
    """Retrieves ISO 9001:2015 clause checklists for the campus (Institutional QMS) per cycle year."""
    target_prog = "GLOBAL"
    target_cycle = cycle_year or "2025 Surveillance"

    existing = db.query(models.ISORequirement).filter(
        models.ISORequirement.program == target_prog,
        models.ISORequirement.cycle_year == target_cycle
    ).all()

    return existing


@app.post("/iso/requirements", response_model=schemas.ISORequirementResponse, status_code=status.HTTP_201_CREATED)
def create_iso_requirement(
    payload: schemas.ISORequirementCreate,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Admin creates a new ISO requirement item for a specific cycle year."""
    new_req = models.ISORequirement(
        program=payload.program,
        iso_clause=payload.iso_clause,
        title=payload.title,
        description=payload.description,
        auditee_office=payload.auditee_office,
        risk_level=payload.risk_level or "Medium",
        status="Not Compliant",
        cycle_year=payload.cycle_year or "2025 Surveillance"
    )
    db.add(new_req)
    db.commit()
    db.refresh(new_req)
    return new_req


@app.put("/iso/requirements/{req_id}", response_model=schemas.ISORequirementResponse)
def update_iso_requirement(
    req_id: str,
    payload: schemas.ISORequirementCreate,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Admin updates an ISO requirement item."""
    req = db.query(models.ISORequirement).filter(models.ISORequirement.id == req_id).first()
    if not req:
        raise HTTPException(status_code=404, detail="ISO requirement not found.")

    req.iso_clause = payload.iso_clause
    req.title = payload.title
    req.description = payload.description
    req.auditee_office = payload.auditee_office
    req.risk_level = payload.risk_level or req.risk_level
    if payload.cycle_year:
        req.cycle_year = payload.cycle_year
    db.commit()
    db.refresh(req)
    return req


@app.delete("/iso/requirements/{req_id}")
def delete_iso_requirement(
    req_id: str,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Admin deletes an ISO requirement item."""
    req = db.query(models.ISORequirement).filter(models.ISORequirement.id == req_id).first()
    if not req:
        raise HTTPException(status_code=404, detail="ISO requirement not found.")
    db.delete(req)
    db.commit()
    return {"message": "ISO requirement deleted successfully."}


@app.post("/iso/upload-evidence")
async def upload_iso_evidence(
    file: UploadFile = File(...),
    requirement_id: str = Form(...),
    document_name: str = Form(...),
    uploaded_by: str = Form(...),
    program: str = Form(...),
    db: Session = Depends(get_db)
):
    """Uploads an evidence file linked to an ISO clause requirement."""
    req = db.query(models.ISORequirement).filter(models.ISORequirement.id == requirement_id).first()
    if not req:
        raise HTTPException(status_code=404, detail="ISO requirement not found.")

    try:
        contents = await file.read()
        safe_filename = file.filename.replace(" ", "_")
        unique_path = f"iso_evidence/{program}/{int(time.time())}_{safe_filename}"

        supabase.storage.from_("documents").upload(
            file=contents,
            path=unique_path,
            file_options={"content-type": file.content_type or "application/pdf"}
        )
        public_url = supabase.storage.from_("documents").get_public_url(unique_path)

        new_evidence = models.ISOEvidence(
            iso_requirement_id=req.id,
            document_name=document_name,
            file_url=public_url,
            uploaded_by=uploaded_by
        )
        db.add(new_evidence)
        
        # Set status to Pending
        req.status = "Pending"
        db.commit()
        db.refresh(req)

        # RAG AI Vector Ingestion
        try:
            extracted_text = ""
            fn_lower = file.filename.lower()
            if fn_lower.endswith(".pdf"):
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(contents))
                for page in pdf_reader.pages:
                    txt = page.extract_text()
                    if txt: extracted_text += txt + "\n"
            elif fn_lower.endswith(".txt"):
                extracted_text = contents.decode("utf-8")
            elif fn_lower.endswith(".docx"):
                import docx
                d = docx.Document(io.BytesIO(contents))
                extracted_text = "\n".join([p.text for p in d.paragraphs])

            if extracted_text.strip():
                vector_store.add_to_vector_db(extracted_text, {
                    "name": document_name,
                    "category": "Accreditation Evidence",
                    "office": req.auditee_office,
                    "program": "GLOBAL",
                    "iso_clause": req.iso_clause,
                    "uploaded_by": uploaded_by,
                    "file_url": public_url
                })
        except Exception as vexc:
            print(f"[upload_iso_evidence] vector store ingestion warning: {vexc}")

        # Audit event
        try:
            supabase.table("system_events_logs").insert({
                "user_email": uploaded_by,
                "event_type": "ISO Evidence Upload",
                "description": f"Uploaded evidence '{document_name}' for {req.iso_clause} ({program})"
            }).execute()
        except Exception:
            pass

        return {"message": "ISO evidence uploaded successfully!", "public_url": public_url}
    except Exception as exc:
        print(f"[upload_iso_evidence] error: {exc}")
        raise HTTPException(status_code=500, detail="Failed to upload ISO evidence file.")


@app.put("/iso/evidence/{evidence_id}/status")
def update_iso_evidence_status(evidence_id: str, status: str = Body(..., embed=True), db: Session = Depends(get_db)):
    """Admin endpoint to update status of an ISO evidence file."""
    ev = db.query(models.ISOEvidence).filter(models.ISOEvidence.id == evidence_id).first()
    if not ev:
        raise HTTPException(status_code=404, detail="Evidence not found")
    ev.status = status
    db.commit()
    return {"message": f"ISO evidence status updated to {status}"}

@app.delete("/iso/evidence/{evidence_id}")
def delete_iso_evidence(
    evidence_id: str,
    db: Session = Depends(get_db)
):
    """Deletes an ISO evidence file."""
    ev = db.query(models.ISOEvidence).filter(models.ISOEvidence.id == evidence_id).first()
    if not ev:
        raise HTTPException(status_code=404, detail="ISO evidence not found.")

    req = ev.requirement
    db.delete(ev)
    db.commit()

    # Re-evaluate requirement status if no evidences remain
    if req and len(req.evidences) == 0:
        req.status = "Not Compliant"
        db.commit()

    return {"message": "ISO evidence removed successfully."}


@app.put("/iso/requirements/{req_id}/status", response_model=schemas.ISORequirementResponse)
def update_iso_status(
    req_id: str,
    payload: schemas.ISOStatusUpdate,
    db: Session = Depends(get_db)
):
    """Updates ISO clause compliance status (e.g. Compliant, Pending, Not Compliant).
    BRIDGING NON-CONFORMITIES TO ACTION PLANS (CAR WORKFLOW):
    When a clause is marked 'Not Compliant', automatically generate a Draft CAR QMS Action Plan (MRC Form 6) assigned to that specific auditee_office!
    """
    req = db.query(models.ISORequirement).filter(models.ISORequirement.id == req_id).first()
    if not req:
        raise HTTPException(status_code=404, detail="ISO requirement not found.")

    req.status = payload.status
    db.commit()
    db.refresh(req)

    # AUTO-GENERATE CAR (MRC Form 6 Action Plan) IF NON-COMPLIANT
    if payload.status == "Not Compliant":
        from datetime import datetime, timedelta
        target_date_str = (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d")
        
        # Check if an action plan for this audit finding already exists
        existing_car = db.query(models.QMSActionPlan).filter(
            models.QMSActionPlan.process_area.like(f"%CAR: {req.iso_clause}%")
        ).first()

        if not existing_car:
            car_plan = models.QMSActionPlan(
                process_area=f"CAR: {req.iso_clause} ({req.title})",
                opportunity_type="Paper",
                opportunity_description=f"Audit Non-Conformity under {req.iso_clause}: {req.description}",
                action_plan=f"Formulate root-cause analysis and corrective measures to close out non-conformity for {req.iso_clause}.",
                personnel_responsible=f"{req.auditee_office} Head / Designated Quality Officer",
                target_date=target_date_str,
                auditee_office=req.auditee_office,
                status="In Progress"
            )
            db.add(car_plan)
            db.commit()

    return req


@app.get("/iso/schedule/{program}", response_model=schemas.IQAScheduleResponse)
def get_iqa_schedule(program: str, db: Session = Depends(get_db)):
    """Retrieves or seeds the dynamic 3-Day IQA Audit Program Schedule for the campus (Institutional QMS)."""
    target_prog = "GLOBAL"
    sched = db.query(models.IQASchedule).filter(models.IQASchedule.program == target_prog).first()
    if not sched:
        sched = models.IQASchedule(
            program=target_prog,
            academic_year="IQA Audit Cycle 2025-2026",
            day1_date="Sept 10, 2025",
            day1_title="Context, Risk & Resource Audit",
            day1_scope="On-site clause audit of Director of Instruction (DOI), College Deans, Financial Management, Property Custodian & SAO. Audit of Clauses 6.1, 7.1 & 8.5.",
            day2_date="Sept 11, 2025",
            day2_title="HR, Data Systems & External Control",
            day2_scope="Audit of HRMO (Clause 7.2), Registrar & MIS (Clause 9.1), Document Controller (Clause 7.5), Library, and BAC Procurement (Clause 8.4).",
            day3_date="Sept 12, 2025",
            day3_title="Consolidation & Closing Meeting",
            day3_scope="Internal data cross-referencing, synthesis of observations, drafting formal audit findings report, and official Closing Ceremony & Certificate Awarding."
        )
        db.add(sched)
        db.commit()
        db.refresh(sched)
    return sched


@app.put("/iso/schedule/{program}", response_model=schemas.IQAScheduleResponse)
def update_iqa_schedule(
    program: str,
    payload: schemas.IQAScheduleUpdate,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Admin updates the 3-Day IQA Audit Program Schedule dates and focus scope for campus QMS."""
    target_prog = "GLOBAL"
    sched = db.query(models.IQASchedule).filter(models.IQASchedule.program == target_prog).first()
    if not sched:
        sched = models.IQASchedule(program=target_prog)
        db.add(sched)

    sched.academic_year = payload.academic_year
    sched.day1_date = payload.day1_date
    sched.day1_title = payload.day1_title
    sched.day1_scope = payload.day1_scope
    sched.day2_date = payload.day2_date
    sched.day2_title = payload.day2_title
    sched.day2_scope = payload.day2_scope
    sched.day3_date = payload.day3_date
    sched.day3_title = payload.day3_title
    sched.day3_scope = payload.day3_scope

    db.commit()
    db.refresh(sched)
    return sched


DEFAULT_IQA_DAYS = [
    {
        "day_number": 1,
        "day_date": "2025-09-10",
        "title": "Context, Risk & Resource Audit",
        "scope": "On-site clause audit of Director of Instruction (DOI), College Deans, Financial Management, Property Custodian & SAO. Audit of Clauses 6.1, 7.1 & 8.5."
    },
    {
        "day_number": 2,
        "day_date": "2025-09-11",
        "title": "HR, Data Systems & External Control",
        "scope": "Audit of HRMO (Clause 7.2), Registrar & MIS (Clause 9.1), Document Controller (Clause 7.5), Library, and BAC Procurement (Clause 8.4)."
    },
    {
        "day_number": 3,
        "day_date": "2025-09-12",
        "title": "Consolidation & Closing Meeting",
        "scope": "Internal data cross-referencing, synthesis of observations, drafting formal audit findings report, and official Closing Ceremony & Certificate Awarding."
    }
]


@app.get("/iso/schedule-days", response_model=List[schemas.IQADayScheduleResponse])
def get_iqa_schedule_days(cycle_year: str = Query("2025 Surveillance"), db: Session = Depends(get_db)):
    """Retrieves dynamic IQA Audit Days for the campus QMS filtered by cycle year."""
    return db.query(models.IQADaySchedule).filter(
        models.IQADaySchedule.program == "GLOBAL",
        models.IQADaySchedule.cycle_year == cycle_year
    ).order_by(models.IQADaySchedule.day_number.asc()).all()


@app.post("/iso/schedule-days", response_model=schemas.IQADayScheduleResponse, status_code=status.HTTP_201_CREATED)
def create_iqa_schedule_day(
    payload: schemas.IQADayScheduleCreate,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Admin creates a new dynamic IQA Audit Day."""
    new_day = models.IQADaySchedule(
        program="GLOBAL",
        cycle_year=payload.cycle_year,
        day_number=payload.day_number,
        day_date=payload.day_date,
        title=payload.title,
        scope=payload.scope
    )
    db.add(new_day)
    db.commit()
    db.refresh(new_day)
    return new_day


@app.put("/iso/schedule-days/{day_id}", response_model=schemas.IQADayScheduleResponse)
def update_iqa_schedule_day(
    day_id: str,
    payload: schemas.IQADayScheduleCreate,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Admin updates an existing dynamic IQA Audit Day."""
    item = db.query(models.IQADaySchedule).filter(models.IQADaySchedule.id == day_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="IQA Audit Day not found.")
    
    item.day_number = payload.day_number
    item.day_date = payload.day_date
    item.title = payload.title
    item.scope = payload.scope

    db.commit()
    db.refresh(item)
    return item


@app.delete("/iso/schedule-days/{day_id}")
def delete_iqa_schedule_day(
    day_id: str,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Admin deletes an IQA Audit Day."""
    item = db.query(models.IQADaySchedule).filter(models.IQADaySchedule.id == day_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="IQA Audit Day not found.")
    db.delete(item)
    db.commit()
    return {"message": "IQA Audit Day deleted successfully."}


# ─────────────────────────────────────────────────────────────────────────────
# QMS OPPORTUNITIES & ACTION PLANS (MRC Form 6)
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/qms/action-plans", response_model=List[schemas.QMSActionPlanResponse])
def get_qms_action_plans(
    cycle_year: str = Query("2025 Surveillance"),
    office: Optional[str] = Query(None),
    db: Session = Depends(get_db)
):
    """Retrieves QMS Action Plans filtered by cycle year and optional auditee office."""
    query = db.query(models.QMSActionPlan).filter(models.QMSActionPlan.cycle_year == cycle_year)
    if office and office != "all":
        query = query.filter(models.QMSActionPlan.auditee_office == office)
    return query.order_by(models.QMSActionPlan.created_at.desc()).all()


@app.post("/qms/action-plans", response_model=schemas.QMSActionPlanResponse, status_code=status.HTTP_201_CREATED)
def create_qms_action_plan(
    payload: schemas.QMSActionPlanCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Creates a new digital QMS Action Plan (MRC Form 6 Opportunities)."""
    new_plan = models.QMSActionPlan(
        cycle_year=payload.cycle_year,
        auditee_office=payload.auditee_office,
        process_area=payload.process_area,
        opportunity_type=payload.opportunity_type,
        opportunity_description=payload.opportunity_description,
        action_plan=payload.action_plan,
        target_date=payload.target_date,
        personnel_responsible=payload.personnel_responsible,
        status=payload.status or "In Progress",
        created_by=current_user.email
    )
    db.add(new_plan)
    db.commit()
    db.refresh(new_plan)
    return new_plan


@app.put("/qms/action-plans/{plan_id}", response_model=schemas.QMSActionPlanResponse)
def update_qms_action_plan(
    plan_id: str,
    payload: schemas.QMSActionPlanUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Updates an existing digital QMS Action Plan."""
    plan = db.query(models.QMSActionPlan).filter(models.QMSActionPlan.id == plan_id).first()
    if not plan:
        raise HTTPException(status_code=404, detail="QMS Action Plan not found.")

    if payload.auditee_office is not None: plan.auditee_office = payload.auditee_office
    if payload.process_area is not None: plan.process_area = payload.process_area
    if payload.opportunity_type is not None: plan.opportunity_type = payload.opportunity_type
    if payload.opportunity_description is not None: plan.opportunity_description = payload.opportunity_description
    if payload.action_plan is not None: plan.action_plan = payload.action_plan
    if payload.target_date is not None: plan.target_date = payload.target_date
    if payload.personnel_responsible is not None: plan.personnel_responsible = payload.personnel_responsible
    if payload.status is not None: 
        plan.status = payload.status
        if payload.status == "Completed" and not plan.actual_completion_date and not payload.actual_completion_date:
            plan.actual_completion_date = datetime.utcnow().strftime("%Y-%m-%d")

    if payload.actual_completion_date is not None: plan.actual_completion_date = payload.actual_completion_date
    if payload.assessment_date is not None: plan.assessment_date = payload.assessment_date
    if payload.assessment_notes is not None: plan.assessment_notes = payload.assessment_notes

    db.commit()
    db.refresh(plan)
    return plan


@app.delete("/qms/action-plans/{plan_id}")
def delete_qms_action_plan(
    plan_id: str,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Deletes a QMS Action Plan."""
    plan = db.query(models.QMSActionPlan).filter(models.QMSActionPlan.id == plan_id).first()
    if not plan:
        raise HTTPException(status_code=404, detail="QMS Action Plan not found.")
    db.delete(plan)
    db.commit()
    return {"message": "QMS Action Plan successfully deleted."}


@app.post("/qms/action-plans/{plan_id}/upload-evidence")
def upload_qms_action_plan_evidence(
    plan_id: str,
    file: UploadFile = File(...),
    document_name: str = Form(...),
    uploaded_by: str = Form("Faculty User"),
    db: Session = Depends(get_db)
):
    """Uploads an evidence file attached directly to a QMS Action Plan."""
    plan = db.query(models.QMSActionPlan).filter(models.QMSActionPlan.id == plan_id).first()
    if not plan:
        raise HTTPException(status_code=404, detail="QMS Action Plan not found.")

    contents = file.file.read()
    file_path = f"qms_evidences/{uuid.uuid4()}_{file.filename}"
    
    file_url = f"http://localhost:8000/uploads/{file_path}"
    try:
        if supabase:
            supabase.storage.from_("repository").upload(file_path, contents)
            file_url = supabase.storage.from_("repository").get_public_url(file_path)
    except Exception as e:
        print(f"Supabase upload notice: {e}")

    evidence = models.QMSEvidence(
        action_plan_id=plan.id,
        document_name=document_name or file.filename,
        file_url=file_url,
        uploaded_by=uploaded_by
    )
    db.add(evidence)
    db.commit()
    db.refresh(evidence)

    return {"message": "Evidence file attached to Action Plan successfully!", "evidence_id": str(evidence.id), "file_url": file_url}


@app.delete("/qms/action-plans/evidence/{evidence_id}")
def delete_qms_evidence(
    evidence_id: str,
    db: Session = Depends(get_db)
):
    """Deletes an evidence file attached to a QMS Action Plan."""
    evidence = db.query(models.QMSEvidence).filter(models.QMSEvidence.id == evidence_id).first()
    if not evidence:
        raise HTTPException(status_code=404, detail="QMS Evidence file not found.")
    db.delete(evidence)
    db.commit()
    return {"message": "Attached evidence file removed."}


# ─────────────────────────────────────────────────────────────────────────────
# PROGRAM ACCREDITATION & HISTORICAL LEVEL UPGRADE
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/accreditation/program/{program_code}", response_model=schemas.ProgramAccreditationResponse)
def get_program_accreditation(program_code: str, db: Session = Depends(get_db)):
    """Fetches the current level and history for a program. Initializes default if none exists."""
    accreditation = db.query(models.ProgramAccreditation).filter(models.ProgramAccreditation.program_code == program_code).first()
    
    if not accreditation:
        # Seed default starting point with initial history
        accreditation = models.ProgramAccreditation(
            program_code=program_code,
            current_level="Candidate Status",
            status="Active"
        )
        db.add(accreditation)
        db.commit()
        db.refresh(accreditation)
        
        # Seed initial baseline milestone
        initial_history = models.AccreditationHistory(
            program_id=accreditation.id,
            level_achieved="Candidate Status",
            remarks="Initial Candidate Status baseline"
        )
        db.add(initial_history)
        db.commit()
        db.refresh(accreditation)
    elif not accreditation.history or len(accreditation.history) == 0:
        # Ensure historical timeline contains at least the current standing
        initial_history = models.AccreditationHistory(
            program_id=accreditation.id,
            level_achieved=accreditation.current_level or "Candidate Status",
            valid_until=accreditation.valid_until,
            remarks=f"Official baseline standing: {accreditation.current_level or 'Candidate Status'}"
        )
        db.add(initial_history)
        db.commit()
        db.refresh(accreditation)
        
    return accreditation


@app.put("/accreditation/program/{program_code}", response_model=schemas.ProgramAccreditationResponse)
def edit_program_accreditation(
    program_code: str, 
    req: schemas.UpgradeAccreditationRequest,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Directly edits the current standing or active evaluation areas and records the milestone in history."""
    accreditation = db.query(models.ProgramAccreditation).filter(models.ProgramAccreditation.program_code == program_code).first()
    if not accreditation:
        accreditation = models.ProgramAccreditation(
            program_code=program_code,
            current_level=req.new_level or "Candidate Status",
            status="Active"
        )
        db.add(accreditation)
        db.commit()
        db.refresh(accreditation)

    parsed_date = None
    if req.valid_until_date:
        try:
            parsed_date = datetime.fromisoformat(req.valid_until_date)
        except Exception:
            pass

    # When standing level is modified, record a milestone entry into timeline history
    if req.new_level:
        accreditation.current_level = req.new_level
        history_log = models.AccreditationHistory(
            program_id=accreditation.id,
            level_achieved=req.new_level,
            valid_until=parsed_date or accreditation.valid_until,
            certificate_url=req.certificate_url,
            remarks=req.remarks or f"Standing calibrated to {req.new_level}"
        )
        db.add(history_log)

    if req.active_areas is not None:
        accreditation.active_areas = req.active_areas
    if parsed_date:
        accreditation.valid_until = parsed_date
    accreditation.updated_at = datetime.utcnow()

    db.commit()
    db.refresh(accreditation)
    return accreditation


@app.post("/accreditation/program/{program_code}/history-milestone", response_model=schemas.ProgramAccreditationResponse)
def add_accreditation_milestone(
    program_code: str,
    req: schemas.UpgradeAccreditationRequest,
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Adds a promotion level milestone to the accreditation timeline with optional certificate."""
    accreditation = db.query(models.ProgramAccreditation).filter(models.ProgramAccreditation.program_code == program_code).first()
    if not accreditation:
        accreditation = models.ProgramAccreditation(
            program_code=program_code,
            current_level=req.new_level or "Candidate Status",
            status="Active"
        )
        db.add(accreditation)
        db.commit()
        db.refresh(accreditation)

    parsed_date = None
    if req.valid_until_date:
        try:
            parsed_date = datetime.fromisoformat(req.valid_until_date)
        except Exception:
            pass

    history_log = models.AccreditationHistory(
        program_id=accreditation.id,
        level_achieved=req.new_level or "Accreditation Update",
        valid_until=parsed_date,
        certificate_url=req.certificate_url,
        remarks=req.remarks or f"Milestone logged for {req.new_level}"
    )
    db.add(history_log)
    db.commit()
    db.refresh(accreditation)
    return accreditation


@app.post("/accreditation/program/{program_code}/upgrade", response_model=schemas.ProgramAccreditationResponse)
def upgrade_program_accreditation(
    program_code: str, 
    req: schemas.UpgradeAccreditationRequest, 
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Officially upgrades a program's level and logs the level achievement into history."""
    accreditation = db.query(models.ProgramAccreditation).filter(models.ProgramAccreditation.program_code == program_code).first()
    if not accreditation:
        accreditation = models.ProgramAccreditation(
            program_code=program_code,
            current_level=req.new_level,
            status="Active"
        )
        db.add(accreditation)
        db.commit()
        db.refresh(accreditation)

    # 1. Log the new achievement to history
    parsed_date = None
    if req.valid_until_date:
        try:
            parsed_date = datetime.fromisoformat(req.valid_until_date)
        except Exception:
            pass

    history_log = models.AccreditationHistory(
        program_id=accreditation.id,
        level_achieved=req.new_level,
        valid_until=parsed_date,
        certificate_url=req.certificate_url,
        remarks=req.remarks or f"Officially promoted to {req.new_level}"
    )
    db.add(history_log)

    # 2. Update the active status
    accreditation.current_level = req.new_level
    if parsed_date:
        accreditation.valid_until = parsed_date
    accreditation.updated_at = datetime.utcnow()

    db.commit()
    db.refresh(accreditation)

    # 3. Add to Audit Log
    try:
        supabase.table("system_events_logs").insert({
            "user_email": admin.email,
            "event_type": "Accreditation Upgrade",
            "description": f"Officially upgraded {program_code} to {req.new_level}"
        }).execute()
    except Exception:
        pass

    try:
        _notify_all_admins(
            db=db,
            n_type="success",
            title=f"AACCUP Level Upgraded: {program_code}",
            message=f"Program {program_code} has been officially upgraded to {req.new_level} by {admin.full_name or admin.email}."
        )
    except Exception:
        pass

    return accreditation


@app.post("/accreditation/history/{history_id}/upload-certificate")
async def upload_historical_certificate(
    history_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Uploads an official certificate file and attaches it to an accreditation history log."""
    history_log = db.query(models.AccreditationHistory).filter(models.AccreditationHistory.id == history_id).first()
    if not history_log:
        raise HTTPException(status_code=404, detail="History record not found.")

    contents = await file.read()
    safe_filename = file.filename.replace(" ", "_")
    unique_path = f"certificates/{int(time.time())}_{safe_filename}"

    try:
        supabase.storage.from_("documents").upload(
            file=contents,
            path=unique_path,
            file_options={"content-type": file.content_type or "application/pdf"}
        )
        public_url = supabase.storage.from_("documents").get_public_url(unique_path)
        
        history_log.certificate_url = public_url
        db.commit()
        db.refresh(history_log)
        return {"message": "Certificate uploaded successfully", "certificate_url": public_url}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to upload certificate: {str(e)}")


# ─────────────────────────────────────────────────────────────────────────────
# AI DOCUMENT GENERATOR (Supports Local Qwen2.5 / Groq)
# ─────────────────────────────────────────────────────────────────────────────

class GenerateDocRequest(BaseModel):
    prompt: str
    temperature: Optional[float] = 0.3

@app.post("/api/generate-document")
@app.post("/generate-document")
def generate_document_endpoint(req: GenerateDocRequest):
    """
    Generates professional document body content using local Qwen2.5 (via Ollama) or Groq.
    Outputs structured markdown formatted for DOCX/PDF export.
    """
    if not req.prompt.strip():
        raise HTTPException(status_code=400, detail="Prompt cannot be empty.")

    system_prompt = (
        "You are a professional document-drafting assistant embedded in an institutional document generator.\n\n"
        "RULES:\n"
        "1. Generate ONLY the body content of the requested document (academic syllabus, business contract, "
        "memorandum, proposal, letter, NDA, certificate, invoice, lesson plan, report, rubric, etc.).\n"
        "2. NEVER generate, describe, or reference a header, footer, letterhead, logo, page-number block, "
        "or company seal. Those are managed client-side.\n"
        "3. Do NOT wrap the response in markdown code fences.\n"
        "4. Format clearly:\n"
        "   - Top-level title as the first line starting with '# ' in Title Case.\n"
        "   - Section headings starting with '## ' in Title Case.\n"
        "   - Use runs of underscores (e.g. '______________') for fill-in blanks (names, dates, amounts, signatures).\n"
        "   - Use '- ' for bulleted lists and standard paragraphs for prose.\n"
        "5. Keep the draft complete, highly professional, and ready for export."
    )

    try:
        response = groq_client.chat.completions.create(
            model="qwen2.5",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": req.prompt.strip()}
            ],
            temperature=req.temperature or 0.3,
            max_tokens=3000
        )
        content = response.choices[0].message.content.strip()
        return {"content": content}
    except Exception as e:
        print(f"[Document Generator Error]: {e}")
        # Try fallback if Groq API key is set
        groq_api_key = os.getenv("GROQ_API_KEY")
        if groq_api_key:
            try:
                from groq import Groq as GroqClient
                g_client = GroqClient(api_key=groq_api_key)
                g_res = g_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": req.prompt.strip()}
                    ],
                    temperature=req.temperature or 0.3,
                    max_tokens=3000
                )
                return {"content": g_res.choices[0].message.content.strip()}
            except Exception as g_err:
                print(f"[Groq Fallback Error]: {g_err}")

        raise HTTPException(
            status_code=500,
            detail=f"Could not generate document with local AI: {str(e)}"
        )


# ─────────────────────────────────────────────────────────────────────────────
# CLIENT SATISFACTION SURVEY (CSS) ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/api/css-responses", response_model=schemas.CssResponseOut)
@app.post("/css-responses", response_model=schemas.CssResponseOut)
def submit_css_response(req: schemas.CssResponseCreate, db: Session = Depends(get_db)):
    """Submits a Client Satisfaction Survey response and saves it to the database."""
    try:
        new_response = models.CssResponse(
            client_type=req.client_type,
            date_of_service=req.date_of_service,
            gender=req.gender,
            age=req.age,
            region=req.region,
            service_availed=req.service_availed,
            campus=req.campus or "Argao",
            office_visited=req.office_visited,
            office_other=req.office_other,
            cc1=req.cc1,
            cc2=req.cc2,
            cc3=req.cc3,
            sqd0=req.sqd0,
            sqd1=req.sqd1,
            sqd2=req.sqd2,
            sqd3=req.sqd3,
            sqd4=req.sqd4,
            sqd5=req.sqd5,
            sqd6=req.sqd6,
            sqd7=req.sqd7,
            sqd8=req.sqd8,
            suggestions=req.suggestions,
            full_name=req.full_name,
            email=req.email
        )
        db.add(new_response)
        db.commit()
        db.refresh(new_response)
        return new_response
    except Exception as e:
        db.rollback()
        # Also try direct Supabase table insertion as graceful fallback
        try:
            payload = req.dict(exclude_unset=True)
            res = supabase.table("css_responses").insert(payload).execute()
            if res.data and len(res.data) > 0:
                return res.data[0]
        except Exception as sb_err:
            print(f"[Supabase CSS Insert Error]: {sb_err}")
        
        raise HTTPException(status_code=500, detail=f"Failed to submit survey response: {str(e)}")


@app.get("/api/css-responses", response_model=List[schemas.CssResponseOut])
@app.get("/css-responses", response_model=List[schemas.CssResponseOut])
def get_css_responses(
    db: Session = Depends(get_db),
    admin: models.User = Depends(get_current_admin)
):
    """Fetches all Client Satisfaction Survey responses (Admin only)."""
    try:
        return db.query(models.CssResponse).order_by(models.CssResponse.created_at.desc()).all()
    except Exception:
        try:
            res = supabase.table("css_responses").select("*").order("created_at", desc=True).execute()
            return res.data or []
        except Exception as sb_err:
            raise HTTPException(status_code=500, detail=f"Failed to retrieve survey responses: {str(sb_err)}")



